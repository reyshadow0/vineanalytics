from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BORDO   = RGBColor(0x9B, 0x1F, 0x42)
DORADO  = RGBColor(0xC9, 0x93, 0x3A)
VERDE   = RGBColor(0x16, 0x7A, 0x4A)
AZUL    = RGBColor(0x1D, 0x4E, 0xD8)
MORADO  = RGBColor(0x6D, 0x28, 0xD9)
OSCURO  = RGBColor(0x11, 0x18, 0x27)
GRIS    = RGBColor(0x6B, 0x72, 0x80)
BLANCO  = RGBColor(0xFF, 0xFF, 0xFF)

HEX = {
    'bordo':  '9B1F42', 'bordo_lt':  'FDF2F5',
    'dorado': 'C9933A', 'dorado_lt': 'FEFCE8',
    'verde':  '167A4A', 'verde_lt':  'F0FDF4',
    'azul':   '1D4ED8', 'azul_lt':   'EFF6FF',
    'morado': '6D28D9', 'morado_lt': 'F5F3FF',
    'gris':   '374151', 'gris_lt':   'F9FAFB',
    'negro':  '111827', 'blanco':    'FFFFFF',
}

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def bg(cell, hex_col):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s  = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  hex_col)
    pr.append(s)

def border_bottom(para, color='C9933A', sz='6'):
    pr  = para._p.get_or_add_pPr()
    bd  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), color)
    bd.append(bot)
    pr.append(bd)

def left_bar(para, color='9B1F42', sz='18'):
    pr = para._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    lf = OxmlElement('w:left')
    lf.set(qn('w:val'), 'single')
    lf.set(qn('w:sz'), sz)
    lf.set(qn('w:space'), '10')
    lf.set(qn('w:color'), color)
    bd.append(lf)
    pr.append(bd)

def h(text, lvl=1, color=BORDO, after=6, before=14):
    sz = {1:20, 2:15, 3:12}[lvl]
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before if lvl==1 else 10)
    p.paragraph_format.space_after  = Pt(after)
    r  = p.add_run(text)
    r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = color
    if lvl == 1: border_bottom(p)
    return p

def p(text='', bold=False, italic=False, sz=11, color=OSCURO,
      align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=5, indent=0):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.space_before = Pt(sb)
    par.paragraph_format.space_after  = Pt(sa)
    par.paragraph_format.left_indent  = Cm(indent)
    if text:
        r = par.add_run(text)
        r.bold=bold; r.italic=italic
        r.font.size=Pt(sz); r.font.color.rgb=color
    return par

def bullet(text, bold_pre='', indent_cm=0.5, color=OSCURO, marker='•'):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.left_indent  = Cm(indent_cm)
    par.paragraph_format.space_after  = Pt(3)
    par.paragraph_format.space_before = Pt(0)
    if bold_pre:
        r1 = par.add_run(bold_pre)
        r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=color
    r2 = par.add_run(text)
    r2.font.size=Pt(11); r2.font.color.rgb=OSCURO

def check(text, ok=True):
    sym = '✔' if ok else '○'
    col = VERDE if ok else GRIS
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.left_indent  = Cm(0.8)
    par.paragraph_format.space_after  = Pt(3)
    r1 = par.add_run(f'{sym}  ')
    r1.font.color.rgb = col; r1.font.size = Pt(11)
    r2 = par.add_run(text)
    r2.font.size = Pt(11); r2.font.color.rgb = OSCURO

def labeled(label, value, label_color=BORDO):
    par = doc.add_paragraph()
    par.paragraph_format.space_after  = Pt(4)
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.left_indent  = Cm(0.3)
    r1 = par.add_run(label + '  ')
    r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=label_color
    r2 = par.add_run(value)
    r2.font.size=Pt(11); r2.font.color.rgb=OSCURO

def tag(cell, text, bg_hex, color=BLANCO, sz=10, bold=True, center=True):
    bg(cell, bg_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = par.add_run(text)
    r.bold=bold; r.font.size=Pt(sz); r.font.color.rgb=color

def flow_table(steps):
    """Tabla numerada de flujo principal."""
    tbl = doc.add_table(rows=len(steps)+1, cols=2)
    tbl.style = 'Table Grid'
    tag(tbl.rows[0].cells[0], 'Paso', HEX['bordo'], sz=10)
    tag(tbl.rows[0].cells[1], 'Acción', HEX['bordo'], sz=10, center=False)
    tbl.rows[0].cells[0].width = Cm(1.5)
    for i, step in enumerate(steps):
        row = tbl.rows[i+1]
        bg(row.cells[0], HEX['bordo_lt'])
        bg(row.cells[1], HEX['blanco'] if i%2==0 else HEX['gris_lt'])
        r0 = row.cells[0].paragraphs[0]
        r0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = r0.add_run(str(i+1))
        rr.bold=True; rr.font.size=Pt(11); rr.font.color.rgb=BORDO
        r1 = row.cells[1].paragraphs[0]
        r1.add_run(step).font.size = Pt(10)
        row.cells[0].width = Cm(1.5)
    doc.add_paragraph()

def ac_table(conditions):
    """Tabla de condiciones de aceptación Given/When/Then."""
    tbl = doc.add_table(rows=len(conditions)+1, cols=3)
    tbl.style = 'Table Grid'
    for j, hdr in enumerate(['Dado (Given)', 'Cuando (When)', 'Entonces (Then)']):
        c_hex = [HEX['verde'], HEX['azul'], HEX['bordo']][j]
        tag(tbl.rows[0].cells[j], hdr, c_hex, sz=10)
    for i, (given, when, then) in enumerate(conditions):
        row = tbl.rows[i+1]
        pairs = [(given, HEX['verde_lt'], VERDE),
                 (when,  HEX['azul_lt'],  AZUL),
                 (then,  HEX['bordo_lt'], BORDO)]
        for j, (txt, bg_h, col) in enumerate(pairs):
            bg(row.cells[j], bg_h)
            r = row.cells[j].paragraphs[0]
            run = r.add_run(txt)
            run.font.size = Pt(9.5); run.font.color.rgb = OSCURO
    doc.add_paragraph()

def uc_header(uc_id, name, actors, priority='Alta', status='Aprobado'):
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = 'Table Grid'
    cells_r0 = tbl.rows[0].cells
    bg(cells_r0[0], HEX['bordo']); bg(cells_r0[1], HEX['bordo_lt'])
    bg(cells_r0[2], HEX['bordo']); bg(cells_r0[3], HEX['bordo_lt'])
    tag(cells_r0[0], 'ID Caso de Uso', HEX['bordo'], sz=9)
    tag(cells_r0[1], uc_id, HEX['bordo_lt'], color=BORDO, bold=True, sz=13)
    tag(cells_r0[2], 'Nombre', HEX['bordo'], sz=9)
    tag(cells_r0[3], name, HEX['bordo_lt'], color=OSCURO, bold=True, sz=11, center=False)
    cells_r1 = tbl.rows[1].cells
    bg(cells_r1[0], HEX['gris_lt']); bg(cells_r1[1], HEX['blanco'])
    bg(cells_r1[2], HEX['gris_lt']); bg(cells_r1[3], HEX['blanco'])
    tag(cells_r1[0], 'Actores', HEX['gris_lt'], color=GRIS, sz=9)
    tag(cells_r1[1], actors, HEX['blanco'], color=OSCURO, bold=False, sz=10, center=False)
    tag(cells_r1[2], 'Prioridad / Estado', HEX['gris_lt'], color=GRIS, sz=9)
    tag(cells_r1[3], f'{priority}  |  {status}', HEX['blanco'], color=OSCURO, bold=False, sz=10, center=False)
    doc.add_paragraph()

def hu_header(hu_id, title, role, priority='Alta', effort='5 pts'):
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = 'Table Grid'
    r0 = tbl.rows[0].cells
    bg(r0[0], HEX['dorado']); bg(r0[1], HEX['dorado_lt'])
    bg(r0[2], HEX['dorado']); bg(r0[3], HEX['dorado_lt'])
    tag(r0[0], 'ID Historia', HEX['dorado'], sz=9)
    tag(r0[1], hu_id, HEX['dorado_lt'], color=DORADO, bold=True, sz=13)
    tag(r0[2], 'Título', HEX['dorado'], sz=9)
    tag(r0[3], title, HEX['dorado_lt'], color=OSCURO, bold=True, sz=11, center=False)
    r1 = tbl.rows[1].cells
    bg(r1[0], HEX['gris_lt']); bg(r1[1], HEX['blanco'])
    bg(r1[2], HEX['gris_lt']); bg(r1[3], HEX['blanco'])
    tag(r1[0], 'Rol / Actor', HEX['gris_lt'], color=GRIS, sz=9)
    tag(r1[1], role, HEX['blanco'], color=OSCURO, bold=False, sz=10, center=False)
    tag(r1[2], 'Prioridad / Esfuerzo', HEX['gris_lt'], color=GRIS, sz=9)
    tag(r1[3], f'{priority}  |  {effort}', HEX['blanco'], color=OSCURO, bold=False, sz=10, center=False)
    doc.add_paragraph()

def story_box(as_role, i_want, so_that):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent  = Cm(0.8)
    par.paragraph_format.right_indent = Cm(0.8)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after  = Pt(10)
    left_bar(par, '9B1F42', '18')
    r1 = par.add_run('Como ')
    r1.font.size=Pt(11); r1.font.color.rgb=GRIS
    r2 = par.add_run(as_role)
    r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=BORDO
    r3 = par.add_run(', quiero ')
    r3.font.size=Pt(11); r3.font.color.rgb=GRIS
    r4 = par.add_run(i_want)
    r4.bold=True; r4.font.size=Pt(11); r4.font.color.rgb=OSCURO
    r5 = par.add_run(', para que ')
    r5.font.size=Pt(11); r5.font.color.rgb=GRIS
    r6 = par.add_run(so_that + '.')
    r6.italic=True; r6.font.size=Pt(11); r6.font.color.rgb=OSCURO

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
pc = doc.add_paragraph()
pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pc.add_run('🍷  VinAnalytics Group')
r.bold=True; r.font.size=Pt(30); r.font.color.rgb=BORDO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Plataforma de Inteligencia Vitivinícola')
r2.italic=True; r2.font.size=Pt(15); r2.font.color.rgb=DORADO

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('ESPECIFICACIÓN DE CASOS DE USO\nE HISTORIAS DE USUARIO')
r3.bold=True; r3.font.size=Pt(16); r3.font.color.rgb=OSCURO

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('Incluye Condiciones de Aceptación (Given / When / Then)')
r4.italic=True; r4.font.size=Pt(11); r4.font.color.rgb=GRIS

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run('Versión 1.0  ·  2025')
r5.font.size=Pt(11); r5.font.color.rgb=GRIS

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
h('1. Introducción', 1)
p('Este documento especifica los Casos de Uso principales y las Historias de Usuario '
  'del sistema VinAnalytics Group, plataforma de Business Intelligence para el análisis '
  'del mercado vitivinícola. Cada caso de uso incluye flujo principal, flujos alternativos '
  'y condiciones de aceptación verificables. Las historias de usuario siguen el formato '
  'estándar ágil con criterios Given/When/Then.')

p('El sistema cuenta con tres roles de usuario:', sa=4)
bullet('Administrador — acceso total al sistema, ETL, usuarios y auditoría.', bold_pre='', color=BORDO)
bullet('Analista — acceso al dashboard, filtros, gráficas y generación de datos.', color=AZUL)
bullet('Gerente — acceso de solo lectura al dashboard y análisis.', color=VERDE)

doc.add_paragraph()
# Tabla resumen de casos de uso
h('Resumen de Casos de Uso', 2, color=OSCURO, before=8)
tbl_res = doc.add_table(rows=8, cols=4)
tbl_res.style = 'Table Grid'
hdrs = ['ID', 'Nombre del Caso de Uso', 'Actor(es)', 'Prioridad']
for j, hdr in enumerate(hdrs):
    tag(tbl_res.rows[0].cells[j], hdr, HEX['bordo'], sz=10)
ucs_summary = [
    ('UC-01', 'Autenticación y Control de Acceso',    'Todos los roles',                'Alta'),
    ('UC-02', 'Ejecutar Pipeline ETL Completo',        'Administrador',                  'Alta'),
    ('UC-03', 'Visualizar Dashboard Analítico',        'Administrador, Analista, Gerente','Alta'),
    ('UC-04', 'Filtrar y Consultar Reseñas de Vino',  'Administrador, Analista',         'Alta'),
    ('UC-05', 'Gestionar Usuarios del Sistema',        'Administrador',                  'Alta'),
    ('UC-06', 'Crear y Gestionar Respaldos',           'Administrador',                  'Media'),
    ('UC-07', 'Generar Datos Aleatorios a StarRocks',  'Administrador, Analista',         'Media'),
]
for i, (uc_id, name, actors, priority) in enumerate(ucs_summary):
    row = tbl_res.rows[i+1]
    bg(row.cells[0], HEX['bordo_lt'])
    bg(row.cells[1], HEX['blanco'] if i%2==0 else HEX['gris_lt'])
    bg(row.cells[2], HEX['blanco'] if i%2==0 else HEX['gris_lt'])
    pri_col = HEX['bordo_lt'] if priority=='Alta' else HEX['dorado_lt']
    bg(row.cells[3], pri_col)
    for j, txt in enumerate([uc_id, name, actors, priority]):
        r = row.cells[j].paragraphs[0]
        run = r.add_run(txt)
        run.font.size=Pt(10)
        if j==0: run.bold=True; run.font.color.rgb=BORDO
        if j==3: run.bold=True; run.font.color.rgb=BORDO if priority=='Alta' else DORADO

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CASOS DE USO
# ══════════════════════════════════════════════════════════════════════════════
h('2. Especificación de Casos de Uso', 1)

# ─────────────────────────────────────────────────────────────────────────────
# UC-01
# ─────────────────────────────────────────────────────────────────────────────
h('UC-01 — Autenticación y Control de Acceso', 2, color=BORDO)
uc_header('UC-01', 'Autenticación y Control de Acceso',
          'Administrador, Analista, Gerente', 'Alta', 'Implementado')

labeled('Descripción:', 'Permite a los usuarios del sistema iniciar sesión con credenciales '
        'válidas y acceder a las funcionalidades según su rol asignado. El sistema '
        'verifica la identidad, establece la sesión y registra el evento en auditoría.')
labeled('Precondiciones:', 'El usuario debe estar registrado en la tabla usuarios_sistema de StarRocks.')
labeled('Postcondiciones:', 'Sesión activa con rol asignado. Evento registrado en auditoría.')

h('Flujo Principal', 3, color=OSCURO, before=6)
flow_table([
    'El usuario accede a localhost:5000 y es redirigido a /login.',
    'El usuario ingresa su nombre de usuario y contraseña en el formulario.',
    'El sistema consulta usuarios_sistema en StarRocks y verifica el hash de la contraseña.',
    'El sistema valida que el usuario esté activo (activo = TRUE).',
    'El sistema crea la sesión Flask con user_id, username y rol.',
    'El sistema registra el evento LOGIN_EXITOSO en la tabla auditoria.',
    'El sistema redirige al usuario al dashboard principal (/).',
])

h('Flujos Alternativos', 3, color=OSCURO, before=4)
bullet('FA-01 Credenciales incorrectas:', ' El sistema muestra "Usuario o contraseña incorrectos" y no crea sesión.')
bullet('FA-02 Usuario inactivo:', ' El sistema muestra "Cuenta desactivada. Contacte al administrador."')
bullet('FA-03 Sesión expirada:', ' El sistema redirige a /login con mensaje de sesión caducada.')

h('Condiciones de Aceptación', 3, color=VERDE, before=6)
ac_table([
    ('El usuario existe en StarRocks\ncon activo=TRUE',
     'El usuario ingresa credenciales\ncorrectas y presiona Iniciar Sesión',
     'El sistema crea sesión, registra\nevento y redirige al dashboard'),
    ('El usuario existe pero\nla contraseña es incorrecta',
     'El usuario ingresa contraseña\nerrónea',
     'Se muestra mensaje de error;\nno se crea sesión; no se registra login'),
    ('El usuario existe pero\nactivo=FALSE',
     'El usuario intenta iniciar sesión',
     'Se muestra mensaje de cuenta\ndesactivada; acceso denegado'),
    ('El usuario no tiene sesión\nactiva',
     'El usuario intenta acceder\na una ruta protegida',
     'El sistema redirige automáticamente\na /login'),
])

# ─────────────────────────────────────────────────────────────────────────────
# UC-02
# ─────────────────────────────────────────────────────────────────────────────
h('UC-02 — Ejecutar Pipeline ETL Completo', 2, color=BORDO)
uc_header('UC-02', 'Ejecutar Pipeline ETL Completo',
          'Administrador', 'Alta', 'Implementado')

labeled('Descripción:', 'El administrador ejecuta las cuatro etapas del pipeline ETL: '
        'carga del CSV a PocketBase, extracción a Parquet, transformación al modelo estrella '
        'y carga final a StarRocks. Cada etapa puede ejecutarse de forma independiente '
        'o en secuencia desde el Panel ETL del dashboard.')
labeled('Precondiciones:', 'Sesión activa con rol Administrador. PocketBase y StarRocks en ejecución. '
        'Archivo winemag-data-130k-v2.csv presente en el directorio del proyecto.')
labeled('Postcondiciones:', 'StarRocks contiene el modelo estrella completo con fact_resenas '
        'y las 6 tablas dimensionales pobladas.')

h('Flujo Principal', 3, color=OSCURO, before=6)
flow_table([
    'El administrador expande el Panel ETL en el dashboard.',
    'El administrador presiona "📤 Cargar CSV → PocketBase". El sistema lee el CSV y sube los registros en lotes de 50.',
    'El sistema confirma la cantidad de registros cargados en el log.',
    'El administrador presiona "Extraer PocketBase → Parquet". El sistema pagina la API de PocketBase y guarda wine_raw.parquet.',
    'El administrador presiona "Transformar y construir dimensiones". El sistema ejecuta transformer.py y genera 7 archivos Parquet.',
    'El administrador selecciona el volumen de carga (Completo / 100k / 50k / 25k) y presiona "Cargar → StarRocks".',
    'El sistema ejecuta loader.py: TRUNCATE + INSERT por lotes en cada tabla del modelo estrella.',
    'El log muestra el resumen final con filas cargadas por tabla.',
])

h('Flujos Alternativos', 3, color=OSCURO, before=4)
bullet('FA-01 PocketBase no disponible:', ' El extractor muestra error de conexión y detiene el proceso.')
bullet('FA-02 CSV no encontrado:', ' pb_loader.py lanza FileNotFoundError con ruta esperada.')
bullet('FA-03 StarRocks no disponible:', ' El loader muestra error de conexión MySQL y hace rollback.')
bullet('FA-04 Carga parcial:', ' El administrador puede usar el selector de registros para cargar subconjuntos.')

h('Condiciones de Aceptación', 3, color=VERDE, before=6)
ac_table([
    ('CSV de 129,971 filas presente,\nPocketBase corriendo',
     'El administrador presiona\n"Cargar CSV → PocketBase"',
     'Los 129,971 registros se cargan\na la colección wine_reviews'),
    ('wine_raw.parquet existe\nen stage/',
     'El administrador presiona\n"Transformar"',
     'Se generan 7 archivos Parquet:\n6 dims + fact_resenas'),
    ('Los 7 parquets existen\nen stage/',
     'El administrador presiona\n"Cargar → StarRocks" (Completo)',
     'fact_resenas contiene 129,971 filas;\n6 dimensiones pobladas'),
    ('StarRocks no está disponible',
     'El sistema intenta conectar\nen la fase de carga',
     'Se muestra error descriptivo en log;\nno se corrompen datos existentes'),
])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# UC-03
# ─────────────────────────────────────────────────────────────────────────────
h('UC-03 — Visualizar Dashboard Analítico', 2, color=BORDO)
uc_header('UC-03', 'Visualizar Dashboard Analítico',
          'Administrador, Analista, Gerente', 'Alta', 'Implementado')

labeled('Descripción:', 'Los usuarios autenticados acceden al dashboard principal que muestra '
        'cinco KPIs en tiempo real, cuatro gráficas analíticas y la tabla de reseñas. '
        'Los datos provienen de consultas directas a StarRocks.')
labeled('Precondiciones:', 'Sesión activa. fact_resenas y tablas dimensionales con datos cargados.')
labeled('Postcondiciones:', 'El usuario visualiza métricas actualizadas del mercado vitivinícola.')

h('Flujo Principal', 3, color=OSCURO, before=6)
flow_table([
    'El usuario accede a / después de autenticarse.',
    'El frontend realiza llamada asíncrona a GET /api/kpis.',
    'StarRocks calcula: COUNT(*) total_resenas, AVG(points), AVG(price), MAX/MIN(price), COUNT dims.',
    'Los 5 KPIs se renderizan: Total Reseñas, Puntuación Promedio, Precio Promedio, Países, Variedades.',
    'El frontend llama a /api/graficas/paises, /variedades, /puntuacion y /bodegas en paralelo.',
    'Chart.js renderiza 4 gráficas: puntuación por país, top variedades, distribución pts, top bodegas.',
    'La tabla llama a GET /api/resenas con paginación (50 registros por página por defecto).',
    'El usuario puede cambiar la paginación y navegar entre páginas.',
])

h('Flujos Alternativos', 3, color=OSCURO, before=4)
bullet('FA-01 Sin datos en StarRocks:', ' Los KPIs muestran 0 y las gráficas muestran "Sin datos — ejecute el pipeline ETL".')
bullet('FA-02 StarRocks no responde:', ' Se muestra "—" en KPIs; la tabla muestra mensaje de error.')

h('Condiciones de Aceptación', 3, color=VERDE, before=6)
ac_table([
    ('fact_resenas contiene\n308,724 registros cargados',
     'El usuario autenticado\naccede al dashboard /',
     'Los 5 KPIs muestran valores\ncorrectos en menos de 3 segundos'),
    ('Los datos están en StarRocks',
     'El dashboard carga las gráficas',
     'Las 4 gráficas se renderizan con\nChart.js sin errores en consola'),
    ('fact_resenas está vacía\n(sin datos)',
     'El usuario accede\nal dashboard',
     'KPIs muestran 0; gráficas muestran\nmensaje "Sin datos — ejecute ETL"'),
    ('El usuario tiene rol Gerente',
     'Intenta acceder a /usuarios\no /auditoria',
     'Es redirigido al dashboard;\nno puede ver páginas de admin'),
])

# ─────────────────────────────────────────────────────────────────────────────
# UC-04
# ─────────────────────────────────────────────────────────────────────────────
h('UC-04 — Filtrar y Consultar Reseñas de Vino', 2, color=BORDO)
uc_header('UC-04', 'Filtrar y Consultar Reseñas de Vino',
          'Administrador, Analista', 'Alta', 'Implementado')

labeled('Descripción:', 'El usuario aplica filtros combinados sobre las 308,724 reseñas '
        'para segmentar el análisis por país, variedad, bodega, puntuación mínima y precio máximo. '
        'Los resultados se paginan y la tabla se actualiza en tiempo real.')
labeled('Precondiciones:', 'Sesión activa. Datos cargados en StarRocks.')
labeled('Postcondiciones:', 'La tabla muestra únicamente los registros que cumplen los criterios seleccionados.')

h('Flujo Principal', 3, color=OSCURO, before=6)
flow_table([
    'El usuario localiza la sección "Filtros de Análisis" en el dashboard.',
    'El usuario selecciona País del desplegable (ej. "France").',
    'Opcionalmente selecciona Variedad (ej. "Pinot Noir").',
    'Opcionalmente ingresa nombre de Bodega (búsqueda parcial LIKE).',
    'Opcionalmente ingresa Puntos mínimos (ej. 92) y Precio máximo (ej. 150).',
    'El usuario presiona "🔍 Aplicar filtros".',
    'El frontend llama a GET /api/resenas con los parámetros activos.',
    'StarRocks ejecuta la consulta con JOINs a las dimensiones y cláusula WHERE dinámica.',
    'La tabla muestra los resultados paginados con el total de registros encontrados.',
])

h('Flujos Alternativos', 3, color=OSCURO, before=4)
bullet('FA-01 Sin resultados:', ' La tabla muestra "Sin registros para los filtros seleccionados."')
bullet('FA-02 Limpiar filtros:', ' El usuario presiona "✕ Limpiar" y se restaura la vista completa.')
bullet('FA-03 Filtro por tecla Enter:', ' Los campos de texto activan el filtro al presionar Enter.')

h('Condiciones de Aceptación', 3, color=VERDE, before=6)
ac_table([
    ('Datos cargados. El usuario\nselecciona País = "Italy"',
     'El usuario presiona\n"Aplicar filtros"',
     'La tabla muestra solo reseñas\nde vinos italianos con total correcto'),
    ('El usuario combina País + Variedad\n+ Puntos mínimos = 95',
     'El usuario presiona\n"Aplicar filtros"',
     'Solo se muestran vinos del país\ny variedad con ≥95 puntos'),
    ('No existen vinos con los\ncriterios seleccionados',
     'El usuario aplica filtros\nmuy restrictivos',
     'La tabla muestra "Sin registros"\ny el contador indica 0 encontrados'),
    ('El usuario ha aplicado filtros',
     'El usuario presiona\n"✕ Limpiar"',
     'Todos los filtros se resetean\ny la tabla muestra todos los registros'),
])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# UC-05
# ─────────────────────────────────────────────────────────────────────────────
h('UC-05 — Gestionar Usuarios del Sistema', 2, color=BORDO)
uc_header('UC-05', 'Gestionar Usuarios del Sistema',
          'Administrador', 'Alta', 'Implementado')

labeled('Descripción:', 'El administrador crea, edita y elimina cuentas de usuario del sistema. '
        'Puede asignar roles (admin/analista/gerente) y activar o desactivar cuentas. '
        'Todas las operaciones quedan registradas en auditoría.')
labeled('Precondiciones:', 'Sesión activa con rol Administrador.')
labeled('Postcondiciones:', 'La tabla usuarios_sistema en StarRocks refleja los cambios. Evento registrado en auditoría.')

h('Flujo Principal — Crear Usuario', 3, color=OSCURO, before=6)
flow_table([
    'El administrador accede a /usuarios.',
    'El administrador completa el formulario: username, contraseña, confirmación y rol.',
    'El sistema valida: username único, contraseña ≥ 6 caracteres, confirmación coincidente.',
    'El sistema genera el hash bcrypt de la contraseña.',
    'El sistema calcula el próximo ID con SELECT COALESCE(MAX(id),0)+1.',
    'El sistema inserta el nuevo usuario en usuarios_sistema.',
    'El sistema registra CREAR_USUARIO en auditoría con IP del solicitante.',
    'Se muestra mensaje de confirmación y la tabla actualizada.',
])

h('Condiciones de Aceptación', 3, color=VERDE, before=6)
ac_table([
    ('El username "analista2" no\nexiste en el sistema',
     'El admin completa el formulario\ncon datos válidos y envía',
     'El usuario se crea en StarRocks;\nse muestra en la tabla; evento en auditoría'),
    ('El username ya existe\nen usuarios_sistema',
     'El admin intenta crear\nun usuario duplicado',
     'Se muestra error "El usuario ya existe";\nno se inserta ningún registro'),
    ('La contraseña tiene\nmenos de 6 caracteres',
     'El admin envía el formulario\ncon contraseña corta',
     'Se muestra validación de error;\nno se procesa la solicitud'),
    ('El admin intenta eliminar\nel usuario con id=1 (admin)',
     'El admin presiona eliminar\nsobre el administrador principal',
     'El sistema bloquea la operación\ny muestra mensaje de protección'),
])

# ─────────────────────────────────────────────────────────────────────────────
# UC-06
# ─────────────────────────────────────────────────────────────────────────────
h('UC-06 — Crear y Gestionar Respaldos', 2, color=BORDO)
uc_header('UC-06', 'Crear y Gestionar Respaldos del Sistema',
          'Administrador', 'Media', 'Implementado')

labeled('Descripción:', 'El administrador genera respaldos JSON del sistema que exportan '
        'los datos de usuarios_sistema y auditoría. Los respaldos se almacenan en la '
        'carpeta backups/ y pueden restaurarse o eliminarse desde la interfaz.')
labeled('Precondiciones:', 'Sesión activa con rol Administrador. StarRocks en ejecución.')
labeled('Postcondiciones:', 'Archivo JSON generado en backups/ con timestamp. Historial de recuperación actualizado.')

h('Flujo Principal', 3, color=OSCURO, before=6)
flow_table([
    'El administrador accede a /respaldos.',
    'El sistema muestra la lista de respaldos existentes con fecha y tamaño.',
    'El administrador presiona "Crear Respaldo Manual".',
    'El sistema consulta usuarios_sistema y auditoría en StarRocks.',
    'El sistema serializa los datos a JSON con timestamp en el nombre del archivo.',
    'El archivo se guarda en backups/backup_YYYYMMDD_HHMMSS.json.',
    'Se actualiza el historial de recuperación y se muestra mensaje de éxito.',
])

h('Condiciones de Aceptación', 3, color=VERDE, before=6)
ac_table([
    ('StarRocks contiene datos en\nusuarios_sistema y auditoria',
     'El admin presiona\n"Crear Respaldo Manual"',
     'Se genera archivo JSON en backups/\ncon timestamp y datos correctos'),
    ('Existen respaldos\npreviamente generados',
     'El admin selecciona un respaldo\ny presiona "Restaurar"',
     'Los datos se restauran en StarRocks;\nse registra evento en auditoría'),
    ('El respaldo automático\nestá configurado (APScheduler)',
     'Pasa el intervalo de tiempo\nprogramado',
     'Se genera respaldo automático\nsin intervención del usuario'),
    ('Existen respaldos en la lista',
     'El admin presiona "Eliminar"\nsobre un respaldo',
     'El archivo se elimina del filesystem;\ndisminuye el contador de respaldos'),
])

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# UC-07
# ─────────────────────────────────────────────────────────────────────────────
h('UC-07 — Generar Datos Aleatorios a StarRocks', 2, color=BORDO)
uc_header('UC-07', 'Generar Datos Aleatorios en StarRocks',
          'Administrador, Analista', 'Media', 'Implementado')

labeled('Descripción:', 'El usuario genera 100,000 reseñas sintéticas con datos aleatorios '
        'e inserta directamente en fact_resenas de StarRocks, usando los IDs de las '
        'dimensiones existentes. Permite simular grandes volúmenes de datos para pruebas analíticas. '
        'Algunos campos se dejan vacíos intencionalmente (precio ~20%, descripción ~12%).')
labeled('Precondiciones:', 'Sesión activa. Las 6 tablas dimensionales deben tener datos (pipeline ETL ejecutado).')
labeled('Postcondiciones:', 'fact_resenas incrementa en 100,000 registros con IDs secuenciales desde MAX(id_resena)+1.')

h('Flujo Principal', 3, color=OSCURO, before=6)
flow_table([
    'El usuario expande el Panel ETL y presiona "🎲 Generar 100k Aleatorios → StarRocks".',
    'El sistema muestra el modal de confirmación con descripción del proceso.',
    'El usuario confirma presionando "🎲 Generar".',
    'El backend consulta los IDs existentes de las 6 tablas dimensionales.',
    'El sistema obtiene el MAX(id_resena) actual para continuar la secuencia.',
    'El sistema genera 100,000 registros aleatorios en lotes de 1,000.',
    'Cada registro usa FKs válidas de las dimensiones existentes.',
    'El sistema hace INSERT en fact_resenas con COMMIT por lote.',
    'El log muestra progreso en tiempo real. Al finalizar, el dashboard se refresca.',
])

h('Condiciones de Aceptación', 3, color=VERDE, before=6)
ac_table([
    ('Las 6 dimensiones tienen datos;\nfact_resenas tiene N registros',
     'El usuario confirma la generación\nde 100,000 registros',
     'fact_resenas pasa a tener N+100,000\nregistros; IDs sin duplicados'),
    ('Las dimensiones están vacías\n(ETL no ejecutado)',
     'El usuario intenta generar\ndatos aleatorios',
     'El sistema muestra error:\n"Ejecute primero el pipeline E→T→L"'),
    ('El proceso de generación\nestá en curso',
     'El usuario presiona otro\nbotón del pipeline',
     'Se muestra "Hay una operación\nen curso" y se bloquea la acción'),
    ('Se generan los 100,000 registros',
     'El proceso finaliza exitosamente',
     'El KPI "Total Reseñas" se actualiza;\nlas gráficas reflejan los nuevos datos'),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# HISTORIAS DE USUARIO
# ══════════════════════════════════════════════════════════════════════════════
h('3. Historias de Usuario', 1)
p('Las historias de usuario describen las funcionalidades desde la perspectiva del usuario '
  'final. Cada historia incluye la narrativa en formato estándar ágil y los criterios de '
  'aceptación verificables en formato Given/When/Then.')

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# HU-01
# ─────────────────────────────────────────────────────────────────────────────
h('HU-01 — Acceso Seguro al Sistema', 2, color=DORADO)
hu_header('HU-01', 'Inicio de Sesión Seguro con Control de Roles',
          'Administrador / Analista / Gerente', 'Alta', '3 pts')
story_box(
    'usuario del sistema',
    'iniciar sesión con mis credenciales y acceder solo a las funciones de mi rol',
    'la información del sistema esté protegida y cada usuario vea únicamente lo que le corresponde'
)
labeled('Criterios de Aceptación:', '', label_color=VERDE)
check('El formulario de login valida campos vacíos antes de enviar.')
check('Las contraseñas se almacenan como hash bcrypt, nunca en texto plano.')
check('Un usuario con rol Gerente NO puede acceder a /usuarios ni /auditoria.')
check('Un login exitoso registra el evento en la tabla auditoria.')
check('Tres intentos fallidos consecutivos no bloquean la cuenta (sin límite implementado).',ok=False)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# HU-02
# ─────────────────────────────────────────────────────────────────────────────
h('HU-02 — Carga Automatizada del Pipeline ETL', 2, color=DORADO)
hu_header('HU-02', 'Ejecutar el Pipeline ETL desde el Dashboard',
          'Administrador', 'Alta', '8 pts')
story_box(
    'administrador del sistema',
    'ejecutar el pipeline completo CSV → PocketBase → Parquet → StarRocks desde el dashboard con un clic por etapa',
    'los datos del mercado vitivinícola estén disponibles para análisis sin necesidad de usar la terminal'
)
labeled('Criterios de Aceptación:', '', label_color=VERDE)
check('Cada etapa del pipeline (CSV, E, T, L) tiene su propio botón en el Panel ETL.')
check('El log de ejecución muestra el progreso en tiempo real con timestamps.')
check('La barra de estado refleja el conteo de registros tras cada etapa.')
check('Si PocketBase no está disponible, el sistema muestra error descriptivo y no bloquea el dashboard.')
check('El selector de límite de carga permite elegir entre carga completa o parcial (25k/50k/100k/completo).')
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# HU-03
# ─────────────────────────────────────────────────────────────────────────────
h('HU-03 — Visualización de KPIs Vitivinícolas', 2, color=DORADO)
hu_header('HU-03', 'Ver Indicadores Clave del Mercado Vitivinícola en Tiempo Real',
          'Gerente / Analista / Administrador', 'Alta', '5 pts')
story_box(
    'gerente de la empresa',
    'ver los indicadores clave del mercado vitivinícola al abrir el dashboard sin necesidad de configuración',
    'pueda tomar decisiones estratégicas basadas en datos actualizados en menos de un minuto'
)
labeled('Criterios de Aceptación:', '', label_color=VERDE)
check('El dashboard muestra 5 KPIs: Total Reseñas, Puntuación Promedio, Precio Promedio, Países y Variedades.')
check('El Precio Promedio calcula solo sobre vinos con price > 0 (excluye sin precio).')
check('El Total Reseñas coincide con el total de la tabla de registros.')
check('Los KPIs se cargan de forma asíncrona sin recargar la página.')
check('Si StarRocks no tiene datos, los KPIs muestran 0 y no lanzan errores en consola.')
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# HU-04
# ─────────────────────────────────────────────────────────────────────────────
h('HU-04 — Análisis de Gráficas del Mercado', 2, color=DORADO)
hu_header('HU-04', 'Visualizar Gráficas Analíticas del Mercado Vitivinícola',
          'Analista / Gerente', 'Alta', '5 pts')
story_box(
    'analista de datos',
    'ver cuatro gráficas interactivas que muestren la puntuación por país, top variedades por precio, distribución de puntuaciones y mejores bodegas',
    'pueda identificar tendencias del mercado y oportunidades de negocio visualmente'
)
labeled('Criterios de Aceptación:', '', label_color=VERDE)
check('La gráfica "Puntuación por País" muestra los 15 países con mayor puntuación promedio (mínimo 80 pts).')
check('La gráfica "Top Variedades" muestra las 12 variedades con mayor precio promedio (excluye price=0).')
check('La gráfica "Distribución de Puntuaciones" muestra el histograma de 80 a 100 puntos.')
check('La gráfica "Top Bodegas" muestra solo bodegas con al menos 10 reseñas.')
check('Al actualizar el dashboard, las gráficas se re-renderizan destruyendo la instancia anterior (sin duplicación).')
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# HU-05
# ─────────────────────────────────────────────────────────────────────────────
h('HU-05 — Filtrado Avanzado de Reseñas', 2, color=DORADO)
hu_header('HU-05', 'Filtrar Reseñas de Vino por Múltiples Criterios',
          'Analista / Administrador', 'Alta', '5 pts')
story_box(
    'analista de datos',
    'filtrar las reseñas de vino combinando país, variedad, bodega, puntuación mínima y precio máximo',
    'pueda segmentar el análisis para identificar los mejores vinos dentro de un presupuesto específico'
)
labeled('Criterios de Aceptación:', '', label_color=VERDE)
check('El filtro por País usa un select con todos los países disponibles en la dimensión dim_pais.')
check('El filtro por Bodega acepta texto parcial (búsqueda LIKE en StarRocks).')
check('Los filtros se aplican con JOIN entre fact_resenas y las tablas dimensionales.')
check('La tabla muestra el total de registros que cumplen los filtros seleccionados.')
check('El botón "Limpiar" resetea todos los filtros y restaura la vista completa.')
check('Los campos numéricos (puntos, precio) activan el filtro al presionar la tecla Enter.')
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# HU-06
# ─────────────────────────────────────────────────────────────────────────────
h('HU-06 — Administración de Usuarios', 2, color=DORADO)
hu_header('HU-06', 'Gestionar Usuarios y Roles del Sistema',
          'Administrador', 'Alta', '8 pts')
story_box(
    'administrador del sistema',
    'crear, editar y desactivar cuentas de usuario asignando roles específicos',
    'pueda controlar quién accede al sistema y con qué nivel de permisos sin necesidad de modificar código'
)
labeled('Criterios de Aceptación:', '', label_color=VERDE)
check('El formulario de creación valida username único antes de intentar insertar en StarRocks.')
check('Las contraseñas se hashean con Werkzeug (pbkdf2:sha256) antes de almacenarse.')
check('Los roles disponibles son: admin, analista, gerente (validados en backend).')
check('El usuario con id=1 (administrador principal) no puede ser eliminado.')
check('Cada operación (crear, editar, eliminar) genera un evento en la tabla auditoria con IP del cliente.')
check('Un usuario desactivado (activo=FALSE) no puede iniciar sesión aunque la contraseña sea correcta.')
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# HU-07
# ─────────────────────────────────────────────────────────────────────────────
h('HU-07 — Generación de Datos de Prueba', 2, color=DORADO)
hu_header('HU-07', 'Generar 100,000 Reseñas Aleatorias para Pruebas Analíticas',
          'Administrador / Analista', 'Media', '3 pts')
story_box(
    'analista de datos',
    'generar 100,000 reseñas sintéticas con datos aleatorios directamente en StarRocks con un clic',
    'pueda probar el comportamiento del dashboard y las gráficas con grandes volúmenes de datos sin depender del CSV original'
)
labeled('Criterios de Aceptación:', '', label_color=VERDE)
check('El sistema solicita confirmación mediante modal personalizado antes de iniciar la generación.')
check('El sistema verifica que las dimensiones existan antes de generar (requiere ETL previo).')
check('Los registros generados usan únicamente IDs de dimensiones existentes (integridad referencial).')
check('Los IDs de fact_resenas continúan desde MAX(id_resena)+1 sin duplicados.')
check('Aproximadamente el 20% de registros tiene price=0 (sin precio), el 12% sin descripción.')
check('El log muestra el progreso en tiempo real (por lotes de 1,000 registros).')
check('Al finalizar, el KPI "Total Reseñas" y las gráficas se actualizan automáticamente.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MATRIZ DE TRAZABILIDAD
# ══════════════════════════════════════════════════════════════════════════════
h('4. Matriz de Trazabilidad — Casos de Uso vs. Historias de Usuario', 1)
p('La siguiente tabla muestra la relación entre cada Historia de Usuario y el '
  'Caso de Uso que implementa.')

tbl_m = doc.add_table(rows=8, cols=4)
tbl_m.style = 'Table Grid'
mhdrs = ['Historia de Usuario', 'Caso de Uso Relacionado', 'Rol Principal', 'Estado']
for j, hdr in enumerate(mhdrs):
    tag(tbl_m.rows[0].cells[j], hdr, HEX['bordo'], sz=10)

matriz = [
    ('HU-01  Acceso Seguro',           'UC-01  Autenticación',        'Todos los roles',  'Implementado'),
    ('HU-02  Pipeline ETL',            'UC-02  Ejecutar ETL',         'Administrador',    'Implementado'),
    ('HU-03  KPIs Vitivinícolas',      'UC-03  Ver Dashboard',        'Gerente/Analista', 'Implementado'),
    ('HU-04  Análisis de Gráficas',    'UC-03  Ver Dashboard',        'Analista',         'Implementado'),
    ('HU-05  Filtrado de Reseñas',     'UC-04  Filtrar Reseñas',      'Analista',         'Implementado'),
    ('HU-06  Gestión de Usuarios',     'UC-05  Gestionar Usuarios',   'Administrador',    'Implementado'),
    ('HU-07  Datos de Prueba',         'UC-07  Generar Aleatorios',   'Analista/Admin',   'Implementado'),
]
for i, (hu, uc, rol, estado) in enumerate(matriz):
    row = tbl_m.rows[i+1]
    bg(row.cells[0], HEX['dorado_lt'])
    bg(row.cells[1], HEX['bordo_lt'])
    bg(row.cells[2], HEX['blanco'] if i%2==0 else HEX['gris_lt'])
    bg(row.cells[3], HEX['verde_lt'])
    for j, (cell, txt) in enumerate(zip(row.cells, [hu, uc, rol, estado])):
        r = cell.paragraphs[0]
        run = r.add_run(txt)
        run.font.size = Pt(10)
        if j==3: run.bold=True; run.font.color.rgb=VERDE

doc.add_paragraph()

# Cierre
p2c = doc.add_paragraph()
p2c.alignment = WD_ALIGN_PARAGRAPH.CENTER
border_bottom(p2c)
border_bottom(p2c, color='9B1F42')
r = p2c.add_run('VinAnalytics Group  ·  Especificación de Casos de Uso e Historias de Usuario  ·  v1.0  ·  2025')
r.italic=True; r.font.size=Pt(10); r.font.color.rgb=GRIS

output = 'VinAnalytics_CasosDeUso_HistoriasUsuario.docx'
doc.save(output)
print(f'[OK] {output}')
