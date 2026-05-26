from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores corporativos ──────────────────────────────────────────────────────
BORDO      = RGBColor(0x9B, 0x1F, 0x42)   # burdeos
DORADO     = RGBColor(0xC9, 0x93, 0x3A)   # dorado
OSCURO     = RGBColor(0x1A, 0x0A, 0x14)   # casi negro
GRIS_CLARO = RGBColor(0xF5, 0xF0, 0xF2)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Márgenes ─────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_paragraph(text='', bold=False, italic=False, size=11,
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p

def add_heading(text, level=1):
    sizes   = {1: 22, 2: 16, 3: 13}
    colors  = {1: BORDO, 2: BORDO, 3: OSCURO}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(sizes[level])
    run.font.color.rgb = colors[level]
    # Línea inferior decorativa para h1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '9B1F42')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_bullet(text, bold_part='', rest='', indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.space_before  = Pt(0)
    if bold_part:
        r = p.add_run(bold_part)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = BORDO
        r2 = p.add_run(rest)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)

def hr():
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9933A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('🍷 VinAnalytics Group')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = BORDO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Plataforma de Inteligencia Vitivinícola')
r2.font.size = Pt(16)
r2.italic = True
r2.font.color.rgb = DORADO

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('DOCUMENTACIÓN ESTRATÉGICA DEL SISTEMA')
r3.bold = True
r3.font.size = Pt(14)
r3.font.color.rgb = OSCURO

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('Los Datos como Activo Estratégico para la Generación de Valor')
r4.font.size = Pt(11)
r4.italic = True
r4.font.color.rgb = RGBColor(0x7A, 0x68, 0x72)

doc.add_paragraph()
doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run('2025')
r5.font.size = Pt(12)
r5.font.color.rgb = RGBColor(0x7A, 0x68, 0x72)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCIÓN DE LA EMPRESA
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('1. Descripción de la Empresa', 1)

add_paragraph(
    'VinAnalytics Group es una empresa de tecnología analítica especializada en '
    'inteligencia de datos para el mercado vitivinícola internacional. Fundada con '
    'el propósito de transformar la manera en que las organizaciones del sector del '
    'vino toman decisiones, la empresa desarrolla y opera plataformas de Business '
    'Intelligence que convierten millones de datos dispersos —reseñas de sumilleres, '
    'precios de mercado, puntuaciones de catadores y tendencias geográficas— en '
    'información estratégica accionable.',
    size=11, space_after=8
)

add_paragraph(
    'La plataforma insignia de la compañía procesa actualmente más de 308,000 reseñas '
    'de vinos provenientes de 44 países, evaluadas por catadores especializados en un '
    'rango de 80 a 100 puntos. Esta base de datos cubre 708 variedades de uva distintas, '
    '16,756 bodegas y 1,230 regiones vitivinícolas, constituyendo uno de los repositorios '
    'analíticos más completos del sector en América Latina.',
    size=11, space_after=8
)

add_paragraph(
    'VinAnalytics Group atiende a distribuidoras de vino, importadoras, cadenas de retail '
    'especializadas, bodegas exportadoras y consultoras del sector, ofreciéndoles '
    'herramientas de análisis que antes solo estaban disponibles para grandes corporaciones '
    'con departamentos de data science internos.',
    size=11, space_after=8
)

add_heading('Datos Clave de la Empresa', 3)
tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
datos = [
    ('Sector', 'Tecnología Analítica / Business Intelligence'),
    ('Especialización', 'Mercado Vitivinícola Internacional'),
    ('Cobertura geográfica', '44 países productores de vino'),
    ('Volumen de datos', '308,724 reseñas · 708 variedades · 16,756 bodegas'),
    ('Tecnología central', 'StarRocks OLAP · Flask · PocketBase · Docker'),
]
for i, (k, v) in enumerate(datos):
    row = tbl.rows[i]
    set_cell_bg(row.cells[0], '1A0A14')
    set_cell_bg(row.cells[1], 'F9F3F5')
    r = row.cells[0].paragraphs[0].add_run(k)
    r.bold = True; r.font.color.rgb = DORADO; r.font.size = Pt(10)
    r2 = row.cells[1].paragraphs[0].add_run(v)
    r2.font.size = Pt(10)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. MISIÓN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('2. Misión', 1)
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Cm(1)
p.paragraph_format.right_indent = Cm(1)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(12)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
for side in ['top','bottom','left','right']:
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), '12' if side == 'left' else '4')
    el.set(qn('w:space'), '4' if side != 'left' else '12')
    el.set(qn('w:color'), '9B1F42' if side == 'left' else 'C9933A')
    pBdr.append(el)
pPr.append(pBdr)
run = p.add_run(
    '"Transformar datos del mercado vitivinícola mundial en inteligencia estratégica '
    'que permita a empresas del sector tomar decisiones informadas, optimizar sus '
    'operaciones comerciales y maximizar su rentabilidad, mediante el desarrollo de '
    'tecnología analítica de vanguardia que convierte el dato en el activo más '
    'valioso de la organización."'
)
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = OSCURO

# ═══════════════════════════════════════════════════════════════════════════════
# 3. VISIÓN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('3. Visión', 1)
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Cm(1)
p.paragraph_format.right_indent = Cm(1)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(12)
pPr = p._p.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
for side in ['top','bottom','left','right']:
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), '12' if side == 'left' else '4')
    el.set(qn('w:space'), '4' if side != 'left' else '12')
    el.set(qn('w:color'), 'C9933A' if side == 'left' else '9B1F42')
    pBdr2.append(el)
pPr.append(pBdr2)
run = p.add_run(
    '"Ser la plataforma de Business Intelligence líder en el sector vitivinícola '
    'de América Latina para el año 2028, reconocida por convertir datos en ventajas '
    'competitivas sostenibles, democratizando el acceso a inteligencia analítica '
    'avanzada para empresas de todos los tamaños dentro de la cadena de valor del vino."'
)
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = OSCURO

# ═══════════════════════════════════════════════════════════════════════════════
# 4. EL DATO COMO ACTIVO ESTRATÉGICO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('4. Los Datos como Activo Estratégico', 1)

add_paragraph(
    'En VinAnalytics Group, los datos no son un subproducto del negocio: son el '
    'núcleo del mismo. Cada reseña procesada, cada puntuación analizada y cada '
    'tendencia de precio detectada representa un activo tangible que genera valor '
    'económico medible para la empresa y sus clientes.',
    size=11, space_after=8
)

add_paragraph(
    'La plataforma transforma datos crudos en tres niveles de valor:', size=11, space_after=4
)
add_bullet('Valor Descriptivo:', ' ¿Qué está pasando en el mercado ahora mismo?')
add_bullet('Valor Diagnóstico:', ' ¿Por qué ciertos vinos o regiones tienen mejor desempeño?')
add_bullet('Valor Predictivo:', ' ¿Qué tendencias emergen y cómo anticiparse a ellas?')

add_paragraph(
    '\nCada uno de estos niveles se traduce directamente en márgenes de ganancia '
    'superiores: reducción de inventario inmovilizado, optimización del portafolio '
    'de productos, identificación de nichos de alto valor y negociación más eficiente '
    'con proveedores y bodegas.',
    size=11, space_after=8
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. TABLA DE OBJETIVOS ESTRATÉGICOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('5. Objetivos Estratégicos e Información para la Toma de Decisiones', 1)

add_paragraph(
    'La siguiente tabla presenta los cinco objetivos estratégicos de VinAnalytics Group, '
    'junto con las necesidades de información asociadas desde la perspectiva táctica '
    '(decisiones de mediano plazo, nivel gerencial) y operativa (decisiones del día a día, '
    'nivel ejecutivo).',
    size=11, space_after=10
)

# Tabla de objetivos
headers = ['Objetivo\nEstratégico', 'Indicadores KPI\nGenerados', 'Nivel Táctico\n(Gerencial)', 'Nivel Operativo\n(Ejecutivo)', 'Impacto en\nMargen']

objetivos = [
    (
        '1. Optimización del\nPortafolio de\nProductos',
        '• Puntuación promedio\n  por variedad\n• Relación precio/\n  calidad\n• Distribución\n  80–100 pts',
        '• Decidir qué variedades\n  priorizar en el\n  portafolio anual\n• Negociar con bodegas\n  basándose en\n  rendimiento histórico\n• Asignar presupuesto\n  de compra por\n  categoría de calidad',
        '• Identificar qué vinos\n  específicos tienen\n  mejor rotación\n• Actualizar catálogo\n  según puntuaciones\n  de catadores\n• Alertas cuando un\n  producto cae en\n  calidad',
        '• Reducción de stock\n  inmovilizado 15–20%\n• Mayor margen en\n  productos premium\n• Menor devolución\n  por calidad'
    ),
    (
        '2. Maximización\ndel Margen\nComercial',
        '• Precio promedio\n  por región/variedad\n• Rango de precios\n  ($4 – $3,300)\n• Precio vs. puntuación\n  (valor percibido)',
        '• Establecer política\n  de precios basada\n  en datos de mercado\n• Identificar productos\n  subvalorados para\n  reposicionamiento\n• Definir segmentos\n  de precio por canal',
        '• Ajuste de precios\n  en tiempo real\n  según competencia\n• Selección de\n  proveedores con\n  mejor relación\n  costo/calidad\n• Negociación de\n  descuentos por\n  volumen informada',
        '• Incremento de\n  margen bruto 8–12%\n• Reducción de\n  compras por debajo\n  del valor de mercado\n• Mayor poder de\n  negociación'
    ),
    (
        '3. Expansión\nGeográfica\nde Mercados',
        '• Reseñas por país\n  (44 países)\n• Calidad promedio\n  por región\n• Tendencias por\n  provincia (426)',
        '• Priorizar qué países\n  o regiones incorporar\n  al portafolio\n• Planificar entrada\n  a nuevos mercados\n  con datos de demanda\n• Evaluar riesgos por\n  concentración\n  geográfica',
        '• Búsqueda diaria de\n  oportunidades en\n  regiones emergentes\n• Monitoreo de\n  puntuaciones de\n  vinos de nuevos\n  países\n• Gestión de\n  proveedores por\n  región',
        '• Diversificación\n  reduce riesgo 25%\n• Acceso a márgenes\n  altos en regiones\n  emergentes\n• Menor dependencia\n  de un solo origen'
    ),
    (
        '4. Fidelización\ny Segmentación\nde Clientes',
        '• Perfil de catadores\n  (20 especialistas)\n• Segmentos por\n  rango de precio\n• Variedades más\n  consultadas',
        '• Diseñar ofertas\n  personalizadas por\n  segmento de cliente\n• Crear programas de\n  fidelización basados\n  en preferencias\n• Definir estrategia\n  de comunicación\n  por perfil',
        '• Recomendar vinos\n  según historial\n  del cliente\n• Alertas de\n  disponibilidad de\n  productos favoritos\n• Atención al cliente\n  informada con datos\n  de calidad',
        '• Aumento del ticket\n  promedio 18–22%\n• Mayor retención\n  de clientes premium\n• Reducción del costo\n  de adquisición'
    ),
    (
        '5. Eficiencia\nOperativa en\nCadena de\nSuministro',
        '• Bodegas con mayor\n  consistencia\n  (16,756 bodegas)\n• Volumen por\n  proveedor\n• Calidad histórica\n  por bodega',
        '• Seleccionar\n  proveedores con\n  historial probado\n  de calidad\n• Planificar compras\n  anuales con datos\n  de tendencias\n• Auditar el\n  desempeño de\n  proveedores',
        '• Control de calidad\n  en recepción basado\n  en benchmarks\n• Optimización de\n  pedidos mínimos\n• Detección temprana\n  de caídas en\n  calidad del proveedor',
        '• Reducción de\n  mermas 10–15%\n• Ahorro en logística\n  por mejor\n  planificación\n• Menor riesgo de\n  quiebre de stock'
    ),
]

# Crear tabla
tbl2 = doc.add_table(rows=len(objetivos)+1, cols=5)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Anchos de columna
widths = [Cm(3.2), Cm(3.0), Cm(3.8), Cm(3.8), Cm(3.0)]
for i, row in enumerate(tbl2.rows):
    for j, cell in enumerate(row.cells):
        cell.width = widths[j]

# Header row
header_texts = ['Objetivo\nEstratégico', 'KPIs del\nSistema', 'Nivel Táctico\n(Gerencial)', 'Nivel Operativo\n(Ejecutivo)', 'Impacto en\nMargen']
for j, (cell, txt) in enumerate(zip(tbl2.rows[0].cells, header_texts)):
    set_cell_bg(cell, '9B1F42')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(txt)
    run.bold = True
    run.font.color.rgb = BLANCO
    run.font.size = Pt(9)

# Data rows
row_colors = ['F9F3F5', 'FFFFFF']
for i, (obj_data) in enumerate(objetivos):
    row = tbl2.rows[i + 1]
    bg = row_colors[i % 2]
    for j, (cell, txt) in enumerate(zip(row.cells, obj_data)):
        if j == 0:
            set_cell_bg(cell, '1A0A14')
        else:
            set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(txt)
        run.font.size = Pt(8.5)
        if j == 0:
            run.bold = True
            run.font.color.rgb = DORADO
        else:
            run.font.color.rgb = OSCURO

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. ALINEACIÓN DE OBJETIVOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('6. Alineación de Objetivos para Máxima Productividad y Margen', 1)

add_paragraph(
    'Los cinco objetivos estratégicos no operan de forma aislada: están diseñados '
    'para potenciarse mutuamente en un ciclo virtuoso que maximiza tanto la '
    'productividad operativa como el margen de ganancia de la empresa.',
    size=11, space_after=10
)

add_heading('6.1 El Ciclo de Valor del Dato', 2)
add_paragraph(
    'La sinergia entre los objetivos se produce a través del siguiente ciclo:', size=11, space_after=6
)

ciclo = [
    ('Captura masiva de datos →',
     'El sistema ingiere más de 308,000 reseñas de 44 países, '
     'creando la base de conocimiento más amplia del sector.'),
    ('Modelado analítico →',
     'El modelo estrella en StarRocks permite consultas sobre millones de '
     'registros en milisegundos, habilitando decisiones en tiempo real.'),
    ('Inteligencia de portafolio →',
     'La información de calidad y precio guía qué comprar, cuánto y a qué '
     'precio vender, optimizando el margen bruto por producto.'),
    ('Expansión geográfica informada →',
     'Los datos de 1,230 regiones revelan oportunidades de mercado antes '
     'invisibles, reduciendo el riesgo de nuevas inversiones.'),
    ('Fidelización basada en datos →',
     'El conocimiento del cliente alimentado por datos de catadores y '
     'variedades genera ventas recurrentes de mayor valor.'),
    ('Eficiencia en proveedores →',
     'El análisis de 16,756 bodegas permite seleccionar proveedores con '
     'mejor relación calidad/precio, reduciendo costos de adquisición.'),
]

for titulo, desc in ciclo:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(titulo + ' ')
    r1.bold = True
    r1.font.color.rgb = BORDO
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

add_heading('6.2 Impacto Cuantificable en el Margen', 2)
add_paragraph(
    'La integración de los cinco objetivos estratégicos genera impactos '
    'medibles en la rentabilidad de la empresa:', size=11, space_after=8
)

impactos = [
    ('Optimización de portafolio',      '+8–12% margen bruto',  'Reducción de productos de bajo rendimiento'),
    ('Política de precios basada en datos', '+15–20% en precio promedio', 'Reposicionamiento de productos subvalorados'),
    ('Expansión geográfica',             'Reducción 25% riesgo', 'Diversificación de proveedores por región'),
    ('Fidelización de clientes',         '+18–22% ticket promedio', 'Clientes premium con recompra frecuente'),
    ('Eficiencia en cadena de suministro', '10–15% menos mermas', 'Proveedores validados con historial de calidad'),
]

tbl3 = doc.add_table(rows=len(impactos)+1, cols=3)
tbl3.style = 'Table Grid'
hdrs3 = ['Objetivo', 'Impacto Esperado', 'Mecanismo']
for j, (cell, txt) in enumerate(zip(tbl3.rows[0].cells, hdrs3)):
    set_cell_bg(cell, 'C9933A')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(txt)
    run.bold = True
    run.font.color.rgb = BLANCO
    run.font.size = Pt(10)

for i, (obj, impacto, mec) in enumerate(impactos):
    row = tbl3.rows[i+1]
    bg = 'F9F3F5' if i % 2 == 0 else 'FFFFFF'
    for j, (cell, txt) in enumerate(zip(row.cells, [obj, impacto, mec])):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(10)
        if j == 1:
            run.bold = True
            run.font.color.rgb = BORDO

doc.add_paragraph()

add_heading('6.3 El Dato como Ventaja Competitiva Sostenible', 2)
add_paragraph(
    'Lo que diferencia a VinAnalytics Group de sus competidores no es solo '
    'tener datos, sino la capacidad de convertirlos en decisiones más rápidas, '
    'más precisas y más rentables que cualquier competidor que opere sin '
    'inteligencia analítica. En un mercado donde la diferencia entre el éxito y '
    'el fracaso se mide en puntos porcentuales de margen, el acceso a información '
    'en tiempo real sobre 44 países, 708 variedades y 16,756 bodegas representa '
    'una barrera competitiva que se hace más fuerte con cada dato nuevo que '
    'ingresa al sistema.',
    size=11, space_after=8
)

add_paragraph(
    'La plataforma está diseñada con esta premisa fundamental: '
    'cada registro almacenado es una decisión futura mejor informada, '
    'y cada decisión mejor informada es un margen de ganancia mayor.',
    bold=True, size=11, color=BORDO, space_after=12
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. ARQUITECTURA TÉCNICA
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('7. Arquitectura Técnica del Sistema', 1)

add_paragraph(
    'La infraestructura tecnológica de VinAnalytics Group está diseñada para '
    'garantizar escalabilidad, disponibilidad y rendimiento analítico de alta velocidad.',
    size=11, space_after=8
)

capas = [
    ('Capa de Ingesta (ETL)',
     'PocketBase actúa como base de staging que recibe los datos crudos del CSV. '
     'El pipeline ETL (Extracción → Transformación → Carga) procesa y normaliza '
     '129,971 registros originales, aplicando lógica de negocio para construir '
     'el modelo estrella.'),
    ('Capa Analítica (StarRocks OLAP)',
     'StarRocks almacena el modelo estrella con 308,724 reseñas y ejecuta consultas '
     'analíticas complejas en milisegundos usando su motor columnar vectorizado. '
     'Soporta crecimiento ilimitado de datos sin degradación de rendimiento.'),
    ('Capa de Presentación (Flask + Chart.js)',
     'El servidor Flask expone una API REST que alimenta el dashboard interactivo. '
     'Las visualizaciones con Chart.js permiten explorar los datos de forma intuitiva '
     'sin conocimientos técnicos.'),
    ('Capa de Seguridad y Auditoría',
     'Sistema de roles (admin/analista/gerente), autenticación segura con hash de '
     'contraseñas, registro completo de auditoría de todas las acciones del sistema '
     'y respaldos automáticos periódicos.'),
    ('Infraestructura (Docker)',
     'Los tres servicios (PocketBase, StarRocks, Flask) corren en contenedores '
     'Docker independientes, garantizando portabilidad, reproducibilidad y '
     'despliegue simplificado en cualquier entorno.'),
]

for titulo, desc in capas:
    add_paragraph(titulo, bold=True, color=BORDO, size=11, space_after=2)
    add_paragraph(desc, size=11, space_after=8)

# ═══════════════════════════════════════════════════════════════════════════════
# CIERRE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
hr()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('VinAnalytics Group — Plataforma de Inteligencia Vitivinícola')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = BORDO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Stack: Flask · StarRocks · PocketBase · Docker · Python 3.13')
r2.font.size = Pt(9)
r2.italic = True
r2.font.color.rgb = RGBColor(0x7A, 0x68, 0x72)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Documento generado en 2025 · Confidencial')
r3.font.size = Pt(9)
r3.font.color.rgb = RGBColor(0x7A, 0x68, 0x72)
hr()

# ── Guardar ───────────────────────────────────────────────────────────────────
output = 'VinAnalytics_Documentacion_Estrategica.docx'
doc.save(output)
print(f'[OK] Documento generado: {output}')
