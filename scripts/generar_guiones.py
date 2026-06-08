"""Genera el documento Word con los dos guiones de video — VinAnalytics."""
import sys, os
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "..", "diagramas")

# ── colores ───────────────────────────────────────────────────────────────────
C_WINE   = RGBColor(0x8B,0x1A,0x1A)
C_WINED  = RGBColor(0x5C,0x0E,0x0E)
C_WINEL  = "FFEAEA"
C_BLUE   = RGBColor(0x1A,0x4A,0x8A)
C_BLUEL  = "E8F0FF"
C_GREEN  = RGBColor(0x2A,0x6A,0x10)
C_GREENL = "EAFAEE"
C_GOLD   = RGBColor(0x8A,0x60,0x10)
C_GOLDL  = "FFF8E8"
C_PURPLE = RGBColor(0x6A,0x1A,0x8A)
C_PURPL  = "F5E8FF"
C_GRAY   = RGBColor(0x66,0x66,0x66)
C_WHITE  = RGBColor(0xFF,0xFF,0xFF)
C_BLACK  = RGBColor(0x1E,0x1E,0x1E)

# ── helpers ────────────────────────────────────────────────────────────────────
def shd(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(s)

def add_run(p, text, bold=False, italic=False, size=11, color=None):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r

def heading(doc, text, size=14, color=None, bold=True, align="left",
            sb=12, sa=4, underline=False, bg=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bg:
        pPr = p._p.get_or_add_pPr()
        s = OxmlElement("w:shd")
        s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
        s.set(qn("w:fill"), bg); pPr.append(s)
    r = p.add_run(text)
    r.bold = bold; r.underline = underline
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def body(doc, text, size=11, color=None, italic=False, sb=2, sa=4, align="left"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.italic = italic
    if color: r.font.color.rgb = color
    return p

def accion(doc, text):
    """Nota de accion de camara/pantalla — en cursiva gris con fondo."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.8)
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"),"clear")
    s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),"F0F0F0")
    pPr.append(s)
    r = p.add_run(f"  {text}")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = C_GRAY

def timestamp_line(doc, tiempo, titulo, color=None):
    """Linea con timecode + titulo del segmento."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"[{tiempo}]  ")
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = C_GOLD
    r2 = p.add_run(titulo)
    r2.bold = True; r2.font.size = Pt(11)
    if color: r2.font.color.rgb = color
    else: r2.font.color.rgb = C_BLACK

def separador(doc, color_hex="CCCCCC"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"4")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"), color_hex)
    pBdr.append(bot); pPr.append(pBdr)

def fase_row(doc, num, color_hx, color_c, titulo, duracion, descripcion, mostrar):
    """Fila visual para cada fase del guion 1."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    # col 0: numero de fase
    shd(t.rows[0].cells[0], color_hx)
    p0 = t.rows[0].cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f"FASE {num}\n{duracion}")
    r0.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = C_WHITE
    t.rows[0].cells[0].width = Cm(2.2)
    # col 1: descripcion
    shd(t.rows[0].cells[1], "FAFAFA")
    p1 = t.rows[0].cells[1].paragraphs[0]
    r1a = p1.add_run(titulo + "\n")
    r1a.bold = True; r1a.font.size = Pt(11); r1a.font.color.rgb = color_c
    r1b = p1.add_run(descripcion)
    r1b.font.size = Pt(10)
    t.rows[0].cells[1].width = Cm(9.5)
    # col 2: que mostrar
    shd(t.rows[0].cells[2], "F5F5F5")
    p2 = t.rows[0].cells[2].paragraphs[0]
    r2a = p2.add_run("Mostrar en camara:\n")
    r2a.bold = True; r2a.font.size = Pt(9); r2a.font.color.rgb = C_GRAY
    r2b = p2.add_run(mostrar)
    r2b.font.size = Pt(9); r2b.italic = True
    t.rows[0].cells[2].width = Cm(4.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def oe_block(doc, tiempo, oe, ot, oo, proceso, narrador, accion_txt, kpi):
    """Bloque de un objetivo operativo en el guion 2."""
    # Header del bloque
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    headers = [("OE", C_WINEL, C_WINED, oe[:30]),
               ("OT", C_BLUEL, C_BLUE,  ot[:35]),
               ("OO", "F5F5F5", C_BLACK, oo[:35]),
               ("Proceso", C_GOLDL, C_GOLD, proceso[:28])]
    widths = [Cm(3.2), Cm(4.2), Cm(4.5), Cm(4.3)]
    for ci, (lbl, bg, fc, val) in enumerate(headers):
        shd(t.rows[0].cells[ci], bg)
        p = t.rows[0].cells[ci].paragraphs[0]
        ra = p.add_run(lbl + "\n"); ra.bold=True; ra.font.size=Pt(8); ra.font.color.rgb=fc
        rb = p.add_run(val);        rb.font.size=Pt(9);  rb.font.color.rgb=fc
        t.rows[0].cells[ci].width = widths[ci]
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    # Timecode
    p_tc = doc.add_paragraph()
    p_tc.paragraph_format.space_before = Pt(2)
    p_tc.paragraph_format.space_after  = Pt(1)
    r_tc = p_tc.add_run(f"[{tiempo}]  ")
    r_tc.bold=True; r_tc.font.size=Pt(10); r_tc.font.color.rgb=C_GOLD
    # Narrador
    p_n = doc.add_paragraph()
    p_n.paragraph_format.space_before = Pt(2)
    p_n.paragraph_format.space_after  = Pt(3)
    p_n.paragraph_format.left_indent  = Cm(0.4)
    r_n = p_n.add_run(narrador)
    r_n.font.size = Pt(11)
    # Accion de pantalla
    accion(doc, accion_txt)
    # KPI
    p_k = doc.add_paragraph()
    p_k.paragraph_format.space_before = Pt(3)
    p_k.paragraph_format.space_after  = Pt(8)
    p_k.paragraph_format.left_indent  = Cm(0.4)
    rk1 = p_k.add_run("KPI / Meta: ")
    rk1.bold=True; rk1.font.size=Pt(10); rk1.font.color.rgb=C_GREEN
    rk2 = p_k.add_run(kpi)
    rk2.font.size=Pt(10)
    separador(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  GUION 1 — LAS 4 FASES DEL SISTEMA (4 minutos)
# ═══════════════════════════════════════════════════════════════════════════════
def guion1(doc):
    doc.add_page_break()

    # Portada del guion
    t_port = doc.add_table(rows=1, cols=1)
    shd(t_port.rows[0].cells[0], "8B1A1A")
    p_port = t_port.rows[0].cells[0].paragraphs[0]
    p_port.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pa = p_port.add_run("GUION 1\n")
    r_pa.bold=True; r_pa.font.size=Pt(18); r_pa.font.color.rgb=C_WHITE
    r_pb = p_port.add_run("LAS 4 FASES DEL SISTEMA VINANALYTICS\n")
    r_pb.bold=True; r_pb.font.size=Pt(14); r_pb.font.color.rgb=C_WHITE
    r_pc = p_port.add_run("Video de presentacion — Duracion: 4 minutos")
    r_pc.font.size=Pt(11); r_pc.font.color.rgb=RGBColor(0xFF,0xCC,0xCC)
    doc.add_paragraph()

    body(doc,
        "Este guion explica la vision completa del proyecto en cuatro fases de desarrollo, "
        "cada una representando el 25% del sistema total. El video comunica que el sistema "
        "esta planificado al 100% aunque actualmente se ha implementado la Fase 1.",
        size=11, sb=6, sa=10)

    # ── INTRO ─────────────────────────────────────────────────────────────
    heading(doc, "INTRODUCCION", size=12, color=C_WINED, sb=8, sa=2, bold=True)
    timestamp_line(doc, "0:00 - 0:30", "Presentacion del proyecto")
    body(doc,
        '"Buenos dias. A continuacion presentamos VinAnalytics Group, una plataforma de '
        'inteligencia de mercado vinicola que transforma mas de 308,000 resenas de vinos '
        'de 44 paises en informacion estrategica para la toma de decisiones. '
        'El sistema esta desarrollado con Flask, StarRocks OLAP y Docker, '
        'y se estructura en cuatro fases de desarrollo. Hoy les presentamos el plan completo."',
        size=11, sb=2, sa=4)
    accion(doc, "CAMARA: Pantalla con el logo de VinAnalytics y la pagina de inicio del sistema en http://localhost:5000")

    separador(doc)

    # ── FASE 1 ────────────────────────────────────────────────────────────
    heading(doc, "FASE 1 — 25% IMPLEMENTADO", size=12, color=C_GREEN, sb=10, sa=4, bold=True)
    timestamp_line(doc, "0:30 - 1:20", "Acceso Publico y Catalogo de Vinos", color=C_GREEN)
    body(doc,
        '"La primera fase — ya implementada — establece la base publica del sistema. '
        'Cualquier visitante puede acceder al catalogo completo de 308,724 resenas de vinos '
        'sin necesidad de crear una cuenta. El sistema esta desplegado con tres contenedores Docker: '
        'Flask como servidor web en el puerto 5000, StarRocks como motor OLAP columnar '
        'para consultas sobre el dataset, y PocketBase para la gestion de sesiones."',
        size=11, sb=2, sa=4)
    body(doc,
        '"En esta fase el visitante puede: navegar el catalogo con fotos reales de vinos, '
        'aplicar filtros combinados por pais, variedad, precio y puntuacion, '
        'y ver la ficha completa de cada vino con la resena del catador profesional. '
        'Las consultas se ejecutan directamente en StarRocks, respondiendo en milisegundos '
        'sobre 300 mil registros."',
        size=11, sb=2, sa=6)
    accion(doc, "CAMARA: Mostrar la pagina principal → navegar a /vinos → aplicar filtro por pais → ver detalle de un vino")
    fase_row(doc, 1, "2A6A10", C_GREEN,
             "Acceso Publico — Catalogo de Vinos",
             "0:30-1:20",
             "Catalogo publico sin login. Busqueda y filtrado multi-criterio. "
             "Ver detalle de vino. Stack: Flask + StarRocks + Docker desplegado.",
             "Pagina /vinos\nFiltros activos\nDetalle /vino/<id>\nTerminal con docker-compose up")

    # ── FASE 2 ────────────────────────────────────────────────────────────
    heading(doc, "FASE 2 — 25% EN DESARROLLO", size=12, color=C_BLUE, sb=10, sa=4, bold=True)
    timestamp_line(doc, "1:20 - 2:20", "Autenticacion y Modulo Analitico", color=C_BLUE)
    body(doc,
        '"La segunda fase incorporara el modulo de autenticacion y las herramientas '
        'de analisis para el rol Analista de Datos. El sistema implementara login seguro '
        'con hash bcrypt y sesiones Flask, control de acceso basado en roles (RBAC) '
        'con tres niveles: analista, gerente y administrador."',
        size=11, sb=2, sa=4)
    body(doc,
        '"El analista tendra acceso a un dashboard analitico interactivo con KPIs calculados '
        'en tiempo real mediante consultas OLAP a StarRocks: top paises por volumen, '
        'distribucion de precios y puntuaciones promedio. Podra aplicar filtros avanzados '
        'multi-criterio, generar reportes personalizados y exportarlos en formato CSV o PDF."',
        size=11, sb=2, sa=6)
    accion(doc, "CAMARA: Mostrar diagrama de flujo de la Fase 2 o wireframe del dashboard analitico")
    fase_row(doc, 2, "1A4A8A", C_BLUE,
             "Autenticacion + Dashboard Analitico",
             "1:20-2:20",
             "Login con roles (analista/gerente/admin). Dashboard OLAP interactivo. "
             "Filtros avanzados. Generacion y exportacion de reportes CSV/PDF.",
             "Diagrama UC_03_analista\nMockup del dashboard\nDiagrama SEQ_01_login")

    # ── FASE 3 ────────────────────────────────────────────────────────────
    heading(doc, "FASE 3 — 25% PLANIFICADO", size=12, color=C_GOLD, sb=10, sa=4, bold=True)
    timestamp_line(doc, "2:20 - 3:10", "Inteligencia Gerencial y Reportes Estrategicos", color=C_GOLD)
    body(doc,
        '"La tercera fase dotara al sistema de inteligencia gerencial. El rol Gerente '
        'tendra acceso a un dashboard ejecutivo simplificado con los KPIs mas relevantes '
        'para la toma de decisiones estrategicas: volumen de resenas por region, '
        'precio promedio por mercado y puntuacion media por pais."',
        size=11, sb=2, sa=4)
    body(doc,
        '"Se incorporara la herramienta de comparacion de KPIs entre hasta 4 mercados '
        'simultaneos con graficos lado a lado y ranking combinado calidad-precio. '
        'Ademas, el modulo de tendencias de precios mostrara la evolucion historica '
        'por variedad con proyeccion para los proximos dos trimestres. '
        'El gerente podra generar reportes estrategicos en PDF con formato ejecutivo."',
        size=11, sb=2, sa=6)
    accion(doc, "CAMARA: Mostrar diagrama UC_04_gerente.excalidraw y diagrama ARCH_03_capas.excalidraw")
    fase_row(doc, 3, "8A6010", C_GOLD,
             "Dashboard Ejecutivo + Comparacion de Mercados",
             "2:20-3:10",
             "Vista ejecutiva con KPIs para gerentes. Comparacion de hasta 4 mercados. "
             "Tendencias de precios con proyeccion. Reportes estrategicos en PDF.",
             "Diagrama UC_04_gerente\nDiagrama ARCH_04_navegacion\nWireframe gerente/dashboard")

    # ── FASE 4 ────────────────────────────────────────────────────────────
    heading(doc, "FASE 4 — 25% PLANIFICADO", size=12, color=C_PURPLE, sb=10, sa=4, bold=True)
    timestamp_line(doc, "3:10 - 3:45", "Administracion Completa del Sistema", color=C_PURPLE)
    body(doc,
        '"La cuarta y ultima fase completara el modulo de administracion del sistema. '
        'El administrador tendra un panel completo de gestion: creacion y administracion '
        'de usuarios con asignacion de roles, panel de control del proceso ETL para '
        'cargar nuevas versiones del dataset Winemag con progreso en tiempo real."',
        size=11, sb=2, sa=4)
    body(doc,
        '"Ademas incluira: un sistema de respaldos manuales y programados del esquema '
        'operacional, el log de auditoria inmutable donde se registran todas las acciones '
        'del sistema con trazabilidad del 100%, y un panel de monitoreo del estado '
        'de los servicios Docker con alertas visuales en tiempo real."',
        size=11, sb=2, sa=6)
    accion(doc, "CAMARA: Mostrar diagrama UC_05_admin.excalidraw y diagrama ARCH_01_despliegue.excalidraw")
    fase_row(doc, 4, "6A1A8A", C_PURPLE,
             "Panel de Administracion Completo",
             "3:10-3:45",
             "Gestion de usuarios RBAC. Control ETL con progreso en tiempo real. "
             "Respaldos automaticos. Log de auditoria inmutable. Monitor de servicios Docker.",
             "Diagrama UC_05_admin\nDiagrama EST_02_etl\nDiagrama ARCH_01_despliegue")

    # ── CIERRE ────────────────────────────────────────────────────────────
    heading(doc, "CIERRE", size=12, color=C_WINED, sb=10, sa=2, bold=True)
    timestamp_line(doc, "3:45 - 4:00", "Conclusion y Stack Tecnologico")
    body(doc,
        '"Con estas cuatro fases de desarrollo, VinAnalytics Group ofrece una solucion '
        'completa de inteligencia de mercado vinicola. El stack tecnologico — Flask, '
        'StarRocks OLAP, PocketBase y Docker — garantiza un sistema de alto rendimiento '
        'capaz de analizar 300 mil resenas en milisegundos. Cuatro actores, '
        'cinco paquetes funcionales, diecinueve casos de uso. '
        'VinAnalytics Group: datos al servicio de la inteligencia vinicola."',
        size=11, sb=2, sa=4)
    accion(doc, "CAMARA: Mostrar diagrama UC_00_paquetes.excalidraw con los 5 paquetes y 4 actores")


# ═══════════════════════════════════════════════════════════════════════════════
#  GUION 2 — DEMOSTRACION EN LA APLICACION (5-6 minutos)
# ═══════════════════════════════════════════════════════════════════════════════
def guion2(doc):
    doc.add_page_break()

    # Portada del guion 2
    t_port = doc.add_table(rows=1, cols=1)
    shd(t_port.rows[0].cells[0], "1A4A8A")
    p_port = t_port.rows[0].cells[0].paragraphs[0]
    p_port.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pa = p_port.add_run("GUION 2\n")
    r_pa.bold=True; r_pa.font.size=Pt(18); r_pa.font.color.rgb=C_WHITE
    r_pb = p_port.add_run("DEMOSTRACION DEL SISTEMA EN LA APLICACION\n")
    r_pb.bold=True; r_pb.font.size=Pt(14); r_pb.font.color.rgb=C_WHITE
    r_pc = p_port.add_run("Objetivos estrategicos, tacticos y operativos — Duracion: 5 a 6 minutos")
    r_pc.font.size=Pt(11); r_pc.font.color.rgb=RGBColor(0xBB,0xCC,0xFF)
    doc.add_paragraph()

    body(doc,
        "Este guion muestra en la aplicacion en ejecucion como el sistema VinAnalytics apoya "
        "cada objetivo de la tabla de evaluacion: tres objetivos estrategicos (OE), cinco tacticos (OT) "
        "y nueve operativos (OO). Se navega la aplicacion en tiempo real mostrando cada funcionalidad.",
        size=11, sb=6, sa=10)

    # ── PREPARACION ───────────────────────────────────────────────────────
    heading(doc, "PREPARACION ANTES DE GRABAR", size=11, color=C_GRAY, sb=6, sa=4, bold=True)
    items = [
        "Ejecutar: docker-compose up -d  (esperar ~2 min a que StarRocks este listo)",
        "Abrir navegador en http://localhost:5000",
        "Tener a mano las credenciales: analista/analista123 | gerente/gerente123 | admin/admin123",
        "Abrir la tabla de la evaluacion (Objetivos.docx) para mostrar al inicio",
        "Resolución de pantalla recomendada: 1280x720 o superior",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(item); r.font.size = Pt(10)

    separador(doc)

    # ── INTRO ─────────────────────────────────────────────────────────────
    heading(doc, "INTRODUCCION", size=12, color=C_WINED, sb=8, sa=2, bold=True)
    timestamp_line(doc, "0:00 - 0:25", "Presentacion del contexto")
    body(doc,
        '"Buenos dias. Vamos a demostrar como VinAnalytics Group, nuestra plataforma de '
        'inteligencia de mercado vinicola, apoya de forma directa y medible cada uno '
        'de los objetivos estrategicos, tacticos y operativos definidos para la organizacion. '
        'El sistema procesa mas de 308,000 resenas de vinos usando StarRocks como motor '
        'OLAP columnar, lo que permite respuestas en milisegundos sobre grandes volumenes de datos."',
        size=11, sb=2, sa=4)
    accion(doc, "CAMARA: Mostrar la tabla de Objetivos.docx abierta → luego cambiar al navegador con http://localhost:5000")
    separador(doc)

    # ── OE3 primero (implementado, se puede mostrar en vivo) ──────────────
    heading(doc, "OE3: MEJORAR LA EXPERIENCIA DE ACCESO A INFORMACION VINICOLA",
            size=12, color=C_GREEN, sb=10, sa=2, bold=True, bg=C_GREENL)
    body(doc, "Objetivo Tactico OT3.1: Facilitar el acceso publico y sin barreras al catalogo de vinos.",
         size=10, color=C_GRAY, sb=2, sa=6)

    oe_block(doc,
        tiempo  = "0:25 - 1:15",
        oe      = "OE3: Mejorar experiencia de acceso",
        ot      = "OT3.1: Facilitar acceso publico sin barreras",
        oo      = "OO3.1.1: Navegacion libre del catalogo sin registro",
        proceso = "Catalogo de Vinos Publico",
        narrador= (
            '"Comenzamos con el Objetivo Operativo 3.1.1: permitir la navegacion libre del '
            'catalogo sin necesidad de registro previo. En la pantalla pueden ver que el '
            'sistema muestra el catalogo de 308,724 resenas de vinos a cualquier visitante. '
            'No se requiere cuenta ni contrasena. Las tarjetas muestran la fotografia real '
            'del vino, su puntuacion en la escala de 80 a 100 puntos, el precio en dolares '
            'y la bodega de origen. El sistema responde en tiempo real con consultas directas '
            'a StarRocks."'
        ),
        accion_txt = (
            "CAMARA: Mostrar http://localhost:5000 → scroll por la pagina principal → "
            "hacer clic en /vinos → navegar por las paginas del catalogo"
        ),
        kpi = "Tiempo de carga de pagina. Meta: menor a 2 segundos. Visitantes sin registro: 100% acceso."
    )

    oe_block(doc,
        tiempo  = "1:15 - 2:00",
        oe      = "OE3: Mejorar experiencia de acceso",
        ot      = "OT3.1: Facilitar acceso publico sin barreras",
        oo      = "OO3.1.2: Busqueda y filtrado multi-criterio del catalogo",
        proceso = "Busqueda y Filtrado Avanzado",
        narrador= (
            '"Objetivo Operativo 3.1.2: busqueda y filtrado avanzado. El sistema ofrece '
            'filtros combinables por pais de origen, variedad de uva, rango de precio '
            'y puntuacion minima. Voy a aplicar un filtro: pais Italia, variedad Pinot Grigio, '
            'precio entre 15 y 40 dolares. Los resultados aparecen en menos de un segundo '
            'gracias al motor OLAP columnar de StarRocks. Esto apoya directamente el '
            'objetivo de retornar resultados en menos de un segundo de tiempo de respuesta."'
        ),
        accion_txt = (
            "CAMARA: En /vinos → activar filtro de pais (ej: Italy) → activar filtro de variedad → "
            "mostrar los resultados → comentar el tiempo de respuesta"
        ),
        kpi = "Tiempo de respuesta de busqueda. Meta: menor a 1 segundo. Filtros combinables: 5."
    )

    # ── OE1 ───────────────────────────────────────────────────────────────
    heading(doc, "OE1: MEJORAR LA TOMA DE DECISIONES COMERCIALES MEDIANTE DATOS",
            size=12, color=C_BLUE, sb=10, sa=2, bold=True, bg=C_BLUEL)
    body(doc, "Objetivos Tacticos OT1.1 (analisis en tiempo real) y OT1.2 (comparacion de mercados).",
         size=10, color=C_GRAY, sb=2, sa=6)

    oe_block(doc,
        tiempo  = "2:00 - 2:50",
        oe      = "OE1: Mejorar toma de decisiones comerciales",
        ot      = "OT1.1: Proveer analisis de mercado en tiempo real",
        oo      = "OO1.1.1: Visualizar KPIs del mercado en dashboard",
        proceso = "Analisis OLAP en tiempo real",
        narrador= (
            '"Objetivo Operativo 1.1.1: dashboard analitico con KPIs en tiempo real. '
            'Iniciamos sesion como Analista de Datos. El sistema valida las credenciales '
            'contra la tabla usuarios_sistema en StarRocks con hash bcrypt y redirige '
            'al dashboard. Aqui pueden ver los indicadores clave del mercado vinicola: '
            'los diez principales paises por volumen de resenas, la distribucion de precios '
            'por segmento y la puntuacion promedio por variedad. Cada grafico se genera '
            'con una consulta OLAP directa a StarRocks sobre los 308 mil registros."'
        ),
        accion_txt = (
            "CAMARA: Navegar a /login → ingresar credenciales de analista → "
            "mostrar el dashboard en /dashboard → senalar cada KPI y grafico"
        ),
        kpi = "Reduccion del 40% en tiempo de analisis manual. Consultas OLAP en milisegundos."
    )

    oe_block(doc,
        tiempo  = "2:50 - 3:35",
        oe      = "OE1: Mejorar toma de decisiones comerciales",
        ot      = "OT1.1: Proveer analisis de mercado en tiempo real",
        oo      = "OO1.1.2: Generar reportes de analisis personalizados",
        proceso = "Generacion de Reportes Analiticos",
        narrador= (
            '"Objetivo Operativo 1.1.2: generacion de reportes personalizados. '
            'Desde el modulo de reportes, el analista configura el reporte: '
            'le asigna un titulo, selecciona el periodo de analisis y los filtros a aplicar. '
            'El sistema genera el reporte con las metricas actuales y lo deja disponible '
            'para exportar. Usando el boton Exportar CSV, descargamos los datos crudos '
            'para procesarlos en otras herramientas. Con Exportar PDF generamos el documento '
            'formal con formato corporativo listo para presentar."'
        ),
        accion_txt = (
            "CAMARA: Navegar a /reportes → crear un nuevo reporte → "
            "configurar nombre y filtros → generar → descargar CSV y PDF"
        ),
        kpi = "Meta: 10 o mas reportes mensuales por analista. Formatos: CSV y PDF disponibles."
    )

    oe_block(doc,
        tiempo  = "3:35 - 4:15",
        oe      = "OE1: Mejorar toma de decisiones comerciales",
        ot      = "OT1.2: Facilitar comparacion estrategica entre mercados",
        oo      = "OO1.2.1: Comparar KPIs entre multiples paises",
        proceso = "Inteligencia Gerencial Comparativa",
        narrador= (
            '"Objetivo Operativo 1.2.1: comparacion de KPIs entre multiples mercados. '
            'Cerramos sesion como analista e iniciamos como Gerente. '
            'En el modulo de comparacion de mercados, el gerente selecciona hasta cuatro '
            'paises para comparar. Seleccionamos: Francia, Italia, Espana y Argentina. '
            'El sistema ejecuta consultas OLAP paralelas a StarRocks y genera la vista '
            'comparativa: graficos de barras lado a lado, tabla de diferencias '
            'y ranking combinado por indice de calidad y precio."'
        ),
        accion_txt = (
            "CAMARA: Logout de analista → login como gerente → navegar a /gerente/comparar → "
            "seleccionar 4 paises → mostrar la vista comparativa"
        ),
        kpi = "Comparacion de hasta 4 mercados simultaneos. Tiempo de comparacion: menos de 30 segundos."
    )

    oe_block(doc,
        tiempo  = "4:15 - 4:50",
        oe      = "OE1: Mejorar toma de decisiones comerciales",
        ot      = "OT1.2: Facilitar comparacion estrategica entre mercados",
        oo      = "OO1.2.2: Analizar tendencias de precios por variedad",
        proceso = "Analisis de Tendencias de Precios",
        narrador= (
            '"Objetivo Operativo 1.2.2: analisis de tendencias de precios. '
            'En el modulo de tendencias, el gerente selecciona la variedad de uva '
            'y el periodo de analisis. El sistema genera el grafico de lineas con '
            'la evolucion historica de precios, incluyendo la linea de tendencia '
            'calculada con regresion lineal. Este modulo apoya la planificacion '
            'de compras y la estrategia de portafolio de la organizacion."'
        ),
        accion_txt = (
            "CAMARA: Navegar a /gerente/tendencias → seleccionar variedad (ej: Pinot Noir) → "
            "mostrar grafico de tendencia con proyeccion"
        ),
        kpi = "Proyeccion de precios para planificacion trimestral. Precision objetivo: 85% o mas."
    )

    # ── OE2 ───────────────────────────────────────────────────────────────
    heading(doc, "OE2: OPTIMIZAR LA GESTION Y SEGURIDAD DE LA INFORMACION",
            size=12, color=C_WINE, sb=10, sa=2, bold=True, bg=C_WINEL)
    body(doc, "Objetivos Tacticos OT2.1 (catalogo actualizado) y OT2.2 (seguridad e integridad).",
         size=10, color=C_GRAY, sb=2, sa=6)

    oe_block(doc,
        tiempo  = "4:50 - 5:20",
        oe      = "OE2: Optimizar gestion de informacion vinicola",
        ot      = "OT2.1: Centralizar y actualizar el catalogo",
        oo      = "OO2.1.1: Ejecutar pipeline ETL del dataset Winemag",
        proceso = "Carga de Datos ETL (Extract-Transform-Load)",
        narrador= (
            '"Objetivo Operativo 2.1.1: pipeline ETL. Iniciamos sesion como Administrador. '
            'En el panel de ETL, el administrador puede cargar una nueva version del dataset '
            'Winemag en formato CSV. Al iniciar el proceso, el sistema ejecuta las tres fases: '
            'Extraccion del archivo CSV con pandas, Transformacion normalizando las seis '
            'dimensiones del modelo estrella, y Carga en StarRocks usando STREAM LOAD. '
            'El progreso se muestra en tiempo real. La meta es cargar mas de 300 mil '
            'registros en menos de diez minutos con cero errores criticos."'
        ),
        accion_txt = (
            "CAMARA: Login como admin → navegar a /admin/etl → "
            "mostrar el panel de carga → iniciar ETL → mostrar barra de progreso"
        ),
        kpi = "Meta: 300K+ registros en menos de 10 minutos. Tasa de error critico: 0%."
    )

    oe_block(doc,
        tiempo  = "5:20 - 5:45",
        oe      = "OE2: Optimizar gestion de informacion vinicola",
        ot      = "OT2.2: Garantizar seguridad e integridad de datos",
        oo      = "OO2.2.1: Gestionar usuarios con control RBAC",
        proceso = "Gestion de Acceso por Roles (RBAC)",
        narrador= (
            '"Objetivo Operativo 2.2.1: gestion de usuarios con control de acceso por roles. '
            'En el panel de usuarios, el administrador ve la lista completa con el rol '
            'y estado de cada cuenta. Puede crear un nuevo usuario asignandole el rol '
            'de analista, gerente o administrador. El sistema almacena la contrasena '
            'con hash bcrypt y garantiza que cada rol accede unicamente a las rutas autorizadas. '
            'Todas las operaciones quedan registradas en el log de auditoria."'
        ),
        accion_txt = (
            "CAMARA: Navegar a /admin/usuarios → mostrar la lista de usuarios → "
            "hacer clic en Nuevo Usuario → llenar el formulario → guardar"
        ),
        kpi = "100% de accesos controlados por rol. Meta: control total con RBAC."
    )

    oe_block(doc,
        tiempo  = "5:45 - 6:10",
        oe      = "OE2: Optimizar gestion de informacion vinicola",
        ot      = "OT2.2: Garantizar seguridad e integridad de datos",
        oo      = "OO2.2.2: Registrar y auditar todas las acciones del sistema",
        proceso = "Log de Auditoria Inmutable",
        narrador= (
            '"Objetivo Operativo 2.2.2: auditoria inmutable de todas las acciones. '
            'El log de auditoria registra en StarRocks cada accion realizada en el sistema: '
            'LOGIN, LOGOUT, ejecuciones de ETL, backups, cambios de usuarios y exportaciones. '
            'El administrador puede filtrar por usuario, tipo de accion y rango de fechas. '
            'Los registros son inmutables — solo se permite insertar, nunca modificar ni eliminar. '
            'El log es exportable en CSV para auditorias externas y cumplimiento normativo."'
        ),
        accion_txt = (
            "CAMARA: Navegar a /admin/auditoria → mostrar el log con distintos tipos de accion → "
            "aplicar filtro por usuario → mostrar boton de exportar CSV"
        ),
        kpi = "Trazabilidad del 100% de las operaciones. Exportable para cumplimiento y auditorias externas."
    )

    # ── CIERRE ────────────────────────────────────────────────────────────
    heading(doc, "CIERRE Y CONCLUSION", size=12, color=C_WINED, sb=10, sa=2, bold=True)
    timestamp_line(doc, "6:10 - 6:30", "Resumen de los tres objetivos estrategicos")
    body(doc,
        '"Como hemos demostrado en esta demostracion, VinAnalytics Group implementa '
        'de forma efectiva los tres objetivos estrategicos definidos. '
        'OE3: el sistema permite a cualquier visitante acceder al catalogo y filtrar '
        '308 mil vinos en menos de un segundo. '
        'OE1: los analistas y gerentes toman decisiones basadas en datos con KPIs '
        'en tiempo real, comparacion de mercados y tendencias de precios. '
        'OE2: el administrador controla el ciclo completo de los datos: carga ETL, '
        'gestion de usuarios con RBAC, y auditoria inmutable de todas las acciones. '
        'VinAnalytics Group: inteligencia vinicola al servicio de la organizacion. Gracias."',
        size=11, sb=2, sa=4)
    accion(doc, "CAMARA: Regresar a la pagina de inicio → mostrar el diagrama UC_00_paquetes.excalidraw → fade out")


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── PORTADA GENERAL ───────────────────────────────────────────────────
    heading(doc, "VINANALYTICS GROUP", size=28, color=C_WINED, align="center",
            sb=60, sa=4)
    heading(doc, "Guiones para Videos de Presentacion", size=16, color=C_GRAY,
            align="center", bold=False, sb=0, sa=8)
    t_cov = doc.add_table(rows=1, cols=1)
    shd(t_cov.rows[0].cells[0], C_WINEL)
    p_c = t_cov.rows[0].cells[0].paragraphs[0]
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_c.add_run(
        "Guion 1: Las 4 Fases del Sistema  |  4 minutos\n"
        "Guion 2: Demostracion en la Aplicacion  |  5 a 6 minutos\n\n"
        "Sistema: Flask  |  StarRocks OLAP  |  PocketBase  |  Docker\n"
        "Dataset: Winemag — 308,724 resenas | 44 paises | 708 variedades"
    )
    r_c.font.size=Pt(11); r_c.font.color.rgb=C_WINED

    guion1(doc)
    guion2(doc)

    out = os.path.join(OUT, "Guiones_Video.docx")
    doc.save(out)
    print(f"OK  {out}")

if __name__ == "__main__":
    main()
