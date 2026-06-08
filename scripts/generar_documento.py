"""Genera el documento Word de documentacion del sistema VinAnalytics."""
import sys, os
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "..", "diagramas")

# ── colores ──────────────────────────────────────────────────────────────────
C_WINE   = RGBColor(0x8B,0x1A,0x1A); C_WINE_HX   = "8B1A1A"
C_WINED  = RGBColor(0x5C,0x0E,0x0E); C_WINED_HX  = "5C0E0E"
C_WINEL_HX = "FFEAEA"
C_BLUE   = RGBColor(0x1A,0x4A,0x8A); C_BLUE_HX   = "1A4A8A"
C_BLUEL_HX = "E8F0FF"
C_GREEN  = RGBColor(0x2A,0x6A,0x10); C_GREEN_HX  = "2A6A10"
C_GREENL_HX= "EAFAEE"
C_PURPLE = RGBColor(0x6A,0x1A,0x8A); C_PURPLE_HX = "6A1A8A"
C_PURPL_HX = "F5E8FF"
C_GOLD_HX  = "8A6010"
C_GOLDL_HX = "FFF8E8"
C_GRAY   = RGBColor(0x55,0x55,0x55); C_GRAY_HX   = "555555"
C_GRAYL_HX = "F5F5F5"
C_WHITE  = RGBColor(0xFF,0xFF,0xFF); C_WHITE_HX  = "FFFFFF"
C_BLACK  = RGBColor(0x1E,0x1E,0x1E)

# ── helpers XML ──────────────────────────────────────────────────────────────
def shd(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(s)

def set_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr")) or tbl.insert(0, OxmlElement("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"CCCCCC")
        tblBorders.append(b)
    tblPr.append(tblBorders)

def para(doc, text, bold=False, size=11, color=None, align="left", space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def heading_styled(doc, text, size=14, color=None, bold=True, align="left",
                   space_before=14, space_after=6, underline=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold; run.underline = underline
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def placeholder_img(doc, description="Screenshot del sistema"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f"[ INSERTAR IMAGEN: {description} ]")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
    r.font.italic = True
    # gray shading on paragraph
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),"EEEEEE")
    pPr.append(s)

# ── SECCION 1: DIAGRAMA DE BASE DE DATOS ────────────────────────────────────
def section_db(doc):
    doc.add_page_break()
    heading_styled(doc, "1. DIAGRAMA DE LA BASE DE DATOS DE LA APLICACION",
                   size=15, color=C_WINED, space_before=0)
    para(doc,
         "VinAnalytics utiliza StarRocks como motor OLAP columnar. El esquema implementa "
         "el modelo estrella (Star Schema): una tabla de hechos central (fact_resenas) "
         "rodeada de seis dimensiones. Ademas existen tablas operacionales para "
         "autenticacion, auditoria y funcionalidades futuras.",
         size=11, space_after=10)

    # ── TABLA FACT ─────────────────────────────────────────────────────────
    heading_styled(doc, "Tabla de Hechos — fact_resenas", size=12,
                   color=C_WINE, space_before=8)
    para(doc, "DUPLICATE KEY(id_resena) | ENGINE=OLAP | DISTRIBUTED BY HASH(id_resena) BUCKETS 10 | 308,724 filas",
         size=9, color=C_GRAY)

    fact_cols = [
        ("id_resena",  "INT",          "PK", "Identificador unico de la resena"),
        ("points",     "INT NOT NULL", "",   "Puntuacion del catador (80-100)"),
        ("price",      "DECIMAL(10,2)","",   "Precio en dolares USD"),
        ("title",      "VARCHAR(500)", "",   "Titulo/nombre completo del vino"),
        ("designation","VARCHAR(300)", "",   "Designacion especifica de la parcela"),
        ("description","VARCHAR(1000)","",   "Texto de la resena del catador"),
        ("region_2",   "VARCHAR(150)", "",   "Sub-region geografica"),
        ("id_pais",    "INT",          "FK", "Referencia a dim_pais"),
        ("id_variedad","INT",          "FK", "Referencia a dim_variedad"),
        ("id_bodega",  "INT",          "FK", "Referencia a dim_bodega"),
        ("id_provincia","INT",         "FK", "Referencia a dim_provincia"),
        ("id_region",  "INT",          "FK", "Referencia a dim_region"),
        ("id_catador", "INT",          "FK", "Referencia a dim_catador"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    set_borders(t)
    t.column_cells(0)[0].width = Cm(3.5)
    t.column_cells(1)[0].width = Cm(3.0)
    t.column_cells(2)[0].width = Cm(1.0)
    t.column_cells(3)[0].width = Cm(7.5)
    hdr = t.rows[0].cells
    for i,(h,bg) in enumerate([("Campo","1A4A8A"),("Tipo","1A4A8A"),("Clave","1A4A8A"),("Descripcion","1A4A8A")]):
        shd(hdr[i], bg)
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = C_WHITE; r.font.size = Pt(10)
    for col,dtype,key,desc in fact_cols:
        row = t.add_row().cells
        bg_col  = "FFF8E8" if key=="PK" else ("FFF0E8" if key=="FK" else C_WHITE_HX)
        key_col = "8A6010" if key=="PK" else ("8B1A1A" if key=="FK" else "333333")
        shd(row[0], bg_col); shd(row[1], bg_col); shd(row[2], bg_col); shd(row[3], bg_col)
        row[0].paragraphs[0].add_run(col).font.size = Pt(10)
        row[1].paragraphs[0].add_run(dtype).font.size = Pt(9)
        r = row[2].paragraphs[0].add_run(key)
        r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(key_col.lstrip("#")) if len(key_col)==6 else C_BLACK
        row[3].paragraphs[0].add_run(desc).font.size = Pt(10)
    doc.add_paragraph()

    # ── TABLAS DIMENSION ───────────────────────────────────────────────────
    heading_styled(doc, "Tablas de Dimension — dim_*", size=12, color=C_BLUE, space_before=8)
    para(doc, "PRIMARY KEY(<pk>) | ENGINE=OLAP | DISTRIBUTED BY HASH(<pk>) BUCKETS 3 | Modelo upsert",
         size=9, color=C_GRAY)

    dims = [
        ("dim_pais",      [("id_pais","INT PK"),("nombre","VARCHAR(100)")],          "44 paises"),
        ("dim_variedad",  [("id_variedad","INT PK"),("nombre","VARCHAR(150)")],       "708 variedades"),
        ("dim_bodega",    [("id_bodega","INT PK"),("nombre","VARCHAR(200)")],         "16,756 bodegas"),
        ("dim_catador",   [("id_catador","INT PK"),("nombre","VARCHAR(150)"),
                           ("twitter","VARCHAR(100)")],                               "Catadores"),
        ("dim_region",    [("id_region","INT PK"),("nombre","VARCHAR(150)")],         "Regiones"),
        ("dim_provincia", [("id_provincia","INT PK"),("nombre","VARCHAR(150)")],      "Provincias"),
    ]
    for tname,cols,nota in dims:
        t2 = doc.add_table(rows=1+len(cols), cols=2)
        set_borders(t2)
        hdr2 = t2.rows[0].cells
        shd(hdr2[0], C_BLUE_HX); shd(hdr2[1], C_BLUE_HX)
        r0 = hdr2[0].paragraphs[0].add_run(tname)
        r0.bold=True; r0.font.color.rgb=C_WHITE; r0.font.size=Pt(10)
        r1 = hdr2[1].paragraphs[0].add_run(nota)
        r1.font.color.rgb=RGBColor(0xBB,0xCC,0xFF); r1.font.size=Pt(9)
        for i,(col,dtype) in enumerate(cols):
            row2 = t2.rows[i+1].cells
            bg2 = C_GOLDL_HX if "PK" in dtype else C_BLUEL_HX
            shd(row2[0], bg2); shd(row2[1], bg2)
            row2[0].paragraphs[0].add_run(col).font.size = Pt(10)
            row2[1].paragraphs[0].add_run(dtype.replace(" PK","")).font.size = Pt(9)
        doc.add_paragraph()

    # ── TABLAS OPERACIONALES ───────────────────────────────────────────────
    heading_styled(doc, "Tablas Operacionales", size=12, color=C_PURPLE, space_before=8)
    ops = [
        ("usuarios_sistema","PRIMARY KEY(id) | Autenticacion y roles",C_WINE_HX,C_WINEL_HX,
         [("id","INT PK"),("username","VARCHAR(50) UNIQUE"),("password_hash","VARCHAR(255)"),
          ("rol","VARCHAR(20) — admin|analista|gerente"),("activo","BOOLEAN"),("created_at","DATETIME")]),
        ("auditoria","DUPLICATE KEY(id) | Log inmutable de acciones",C_PURPLE_HX,C_PURPL_HX,
         [("id","INT PK"),("usuario","VARCHAR(50)"),("rol","VARCHAR(20)"),
          ("accion","VARCHAR(100) — LOGIN|ETL|BACKUP|..."),("detalle","VARCHAR(1000)"),
          ("ip","VARCHAR(50)"),("fecha","DATETIME")]),
        ("favoritos (Fase 3)","PRIMARY KEY(id) | Vinos guardados por usuario",C_GREEN_HX,C_GREENL_HX,
         [("id","INT PK"),("user_id","INT FK"),("id_resena","INT FK"),("created_at","DATETIME")]),
    ]
    for tname,nota,hbg,bbg,cols in ops:
        t3 = doc.add_table(rows=1+len(cols), cols=2)
        set_borders(t3)
        h3 = t3.rows[0].cells
        shd(h3[0], hbg); shd(h3[1], hbg)
        r0=h3[0].paragraphs[0].add_run(tname); r0.bold=True; r0.font.color.rgb=C_WHITE; r0.font.size=Pt(10)
        r1=h3[1].paragraphs[0].add_run(nota); r1.font.color.rgb=C_WHITE; r1.font.size=Pt(9)
        for i,(col,dtype) in enumerate(cols):
            row3 = t3.rows[i+1].cells
            shd(row3[0], bbg); shd(row3[1], bbg)
            row3[0].paragraphs[0].add_run(col).font.size = Pt(10)
            row3[1].paragraphs[0].add_run(dtype).font.size = Pt(9)
        doc.add_paragraph()

    para(doc,
         "NOTA: StarRocks usa ENGINE=OLAP con replication_num=1 (desarrollo). "
         "Las dimensiones usan modelo PRIMARY KEY para soporte de upsert. "
         "La tabla de hechos usa modelo DUPLICATE KEY para maxima velocidad de ingesta.",
         size=9, color=C_GRAY, space_before=4)


# ── SECCION 2: DIAGRAMA DE CASOS DE USO ─────────────────────────────────────
def section_uc_diagram(doc):
    doc.add_page_break()
    heading_styled(doc, "2. DIAGRAMA DE CASOS DE USO DEL SISTEMA",
                   size=15, color=C_WINED, space_before=0)
    para(doc,
         "El sistema VinAnalytics cuenta con cuatro actores principales: Visitante Publico "
         "(sin autenticacion), Analista de Datos, Gerente y Administrador. "
         "Los diagramas completos en formato Excalidraw se encuentran en los archivos "
         "UC_01_general.excalidraw (vision global) y UC_02 a UC_05 (por cada actor). "
         "A continuacion se presenta la matriz actor-caso de uso.",
         size=11, space_after=10)

    placeholder_img(doc, "UC_01_general.excalidraw — abrir en excalidraw.com")

    doc.add_paragraph()
    heading_styled(doc, "Matriz Actor — Caso de Uso", size=12, color=C_BLUE, space_before=8)

    matrix = [
        ("CU-01","Explorar Catalogo de Vinos",        True, True, True, True),
        ("CU-02","Buscar y Filtrar Vinos",             True, True, True, True),
        ("CU-03","Ver Detalle de Vino",                True, True, True, True),
        ("CU-04","Consultar Estadisticas del Catalogo",True, True, True, True),
        ("CU-05","Iniciar Sesion",                     False,True, True, True),
        ("CU-06","Cerrar Sesion",                      False,True, True, True),
        ("CU-07","Visualizar Dashboard Analitico",     False,True, False,False),
        ("CU-08","Aplicar Filtros Avanzados",          False,True, False,False),
        ("CU-09","Generar Reporte de Analisis",        False,True, False,False),
        ("CU-10","Exportar Datos CSV/PDF",             False,True, False,False),
        ("CU-11","Ver Dashboard Ejecutivo",            False,False,True, False),
        ("CU-12","Comparar KPIs por Mercado",          False,False,True, False),
        ("CU-13","Generar Reporte Estrategico",        False,False,True, False),
        ("CU-14","Analizar Tendencias de Precios",     False,False,True, False),
        ("CU-15","Gestionar Usuarios del Sistema",     False,False,False,True),
        ("CU-16","Ejecutar Proceso ETL",               False,False,False,True),
        ("CU-17","Gestionar Respaldos",                False,False,False,True),
        ("CU-18","Consultar Log de Auditoria",         False,False,False,True),
        ("CU-19","Monitorear Estado del Sistema",      False,False,False,True),
    ]
    t = doc.add_table(rows=1+len(matrix), cols=6)
    set_borders(t)
    headers = ["ID","Caso de Uso","Visitante","Analista","Gerente","Admin"]
    hbgs    = [C_GRAY_HX,C_GRAY_HX,C_GREEN_HX,C_BLUE_HX,C_GOLD_HX,C_PURPLE_HX]
    for i,(h,bg) in enumerate(zip(headers,hbgs)):
        c = t.rows[0].cells[i]; shd(c, bg)
        r = c.paragraphs[0].add_run(h)
        r.bold=True; r.font.color.rgb=C_WHITE; r.font.size=Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_colors = [C_GREENL_HX]*4 + [C_WINEL_HX]*2 + [C_BLUEL_HX]*4 + [C_GOLDL_HX]*4 + [C_PURPL_HX]*5
    for ri,(cid,name,vis,ana,ger,adm) in enumerate(matrix):
        row = t.rows[ri+1].cells
        bg = row_colors[ri]
        for c in row: shd(c, bg)
        row[0].paragraphs[0].add_run(cid).font.size = Pt(9)
        row[1].paragraphs[0].add_run(name).font.size = Pt(10)
        for ci,val in enumerate([vis,ana,ger,adm]):
            mark = "SI" if val else "-"
            r = row[ci+2].paragraphs[0].add_run(mark)
            r.font.size = Pt(10); r.bold = val
            if val: r.font.color.rgb = C_GREEN
            row[ci+2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ── SECCION 3: DESCRIPCION DE CASOS DE USO ──────────────────────────────────
PACKAGES = [
  { "num":1,"name":"ACCESO PUBLICO","hbg":C_GREEN_HX,"fgc":C_GREEN,
    "desc":"Funcionalidades disponibles para cualquier visitante sin necesidad de autenticacion. "
           "Representan el punto de entrada del sistema y la propuesta de valor publica.",
    "cases":[
      { "id":"CU-01","name":"Explorar Catalogo de Vinos","priority":9,"scope":"OPERATIVO",
        "purpose":"Permitir que cualquier visitante navegue el catalogo completo de mas de 308,000 "
                   "resenas de vinos sin requerir autenticacion ni registro previo.",
        "desc":"El sistema presenta una galeria paginada de vinos con imagenes reales, puntuacion "
               "(80-100 puntos), precio en USD, bodega y variedad de uva. Las tarjetas se renderizan "
               "con fotos de alta calidad de Unsplash CDN. El usuario puede navegar libremente por "
               "paginas sin restricciones de acceso. La paginacion esta implementada con parametros "
               "GET (?page=N) y retorna 20 vinos por pagina.",
        "stories":[
          ("Visitante","explorar el catalogo completo de vinos disponibles","descubrir nuevas etiquetas sin necesidad de registrarme"),
          ("Aficionado al vino","ver fotos y puntuaciones de diferentes vinos","elegir mi proxima compra con informacion confiable"),
          ("Estudiante de enologia","navegar libremente por todas las resenas del dataset","estudiar variedades, bodegas y regiones vinicolas del mundo"),
          ("Distribuidor comercial","revisar el listado completo del catalogo de vinos","identificar etiquetas de interes para mi portafolio de distribucion"),
        ],
        "img":"Pagina /vinos — galeria de tarjetas con fotos reales de vinos, puntuaciones y precios"
      },
      { "id":"CU-02","name":"Buscar y Filtrar Vinos","priority":8,"scope":"OPERATIVO",
        "purpose":"Proporcionar mecanismos de busqueda y filtrado multi-criterio para que cualquier "
                   "usuario pueda encontrar vinos especificos segun sus preferencias.",
        "desc":"El sistema ofrece filtros combinables por: pais de origen, variedad de uva, rango "
               "de precio (min-max) y rango de puntuacion. La busqueda por texto libre aplica sobre "
               "el titulo y la descripcion del vino. Los resultados se paginan y pueden ordenarse "
               "por precio, puntuacion o relevancia. Los filtros se envian como parametros GET.",
        "stories":[
          ("Visitante","filtrar vinos por pais de origen","explorar la produccion vinicola de una region geografica especifica"),
          ("Comprador","buscar vinos dentro de un rango de precio determinado","encontrar opciones que se ajusten exactamente a mi presupuesto"),
          ("Sommelier","filtrar por variedad de uva especifica","recomendar vinos de una cepa concreta a los clientes de mi restaurante"),
          ("Gerente de restaurante","buscar vinos con puntuacion mayor a 90 puntos","seleccionar los mejores vinos para incluir en la carta del establecimiento"),
        ],
        "img":"Pagina /buscar — formulario de filtros con resultados paginados"
      },
      { "id":"CU-03","name":"Ver Detalle de Vino","priority":7,"scope":"OPERATIVO",
        "purpose":"Mostrar la ficha completa de un vino incluyendo descripcion del catador, "
                   "procedencia geografica y caracteristicas del producto.",
        "desc":"Al seleccionar un vino del catalogo, el sistema muestra la ficha completa: titulo, "
               "bodega, variedad, denominacion, region, pais, puntuacion, precio y descripcion "
               "del catador con su nombre/twitter. La pagina sugiere vinos similares (mismo pais "
               "o variedad) generando una consulta a StarRocks. La URL es /vino/<id_resena>.",
        "stories":[
          ("Visitante","ver la descripcion completa de un vino especifico","conocer su perfil de sabor y caracteristicas antes de comprarlo"),
          ("Enofilo","ver los datos del catador que escribio la resena","evaluar la credibilidad y el perfil del critico que califico el vino"),
          ("Importador","acceder a la region y denominacion exacta del vino","verificar la procedencia geografica y autenticidad del producto"),
          ("Chef","leer las notas de cata detalladas","identificar los aromas y sabores para sugerir maridajes apropiados en el menu"),
        ],
        "img":"Pagina /vino/<id> — ficha completa del vino con todos sus atributos"
      },
      { "id":"CU-04","name":"Consultar Estadisticas del Catalogo","priority":5,"scope":"TACTICO",
        "purpose":"Ofrecer al visitante una vista de alto nivel con estadisticas globales del "
                   "catalogo: distribucion por pais, rangos de precio y puntuacion promedio.",
        "desc":"La pagina de inicio presenta estadisticas agregadas calculadas mediante consultas "
               "OLAP a StarRocks: total de resenas (308,724), cantidad de paises (44), precio "
               "promedio, puntuacion media y top 5 paises por volumen. Los datos se consultan "
               "en tiempo real y se presentan en tarjetas de categorias con imagen de fondo.",
        "stories":[
          ("Investigador","ver estadisticas generales del mercado vinicola mundial","entender la distribucion de la produccion por regiones geograficas"),
          ("Estudiante universitario","consultar cuantos vinos existen por pais en el dataset","obtener datos cuantitativos para mi tesis sobre el mercado vinicola global"),
          ("Periodista","acceder al precio promedio por region","redactar un articulo bien fundamentado sobre tendencias del mercado del vino"),
          ("Inversor","conocer la distribucion de puntuaciones del catalogo completo","evaluar la calidad promedio del mercado antes de tomar decisiones de inversion"),
        ],
        "img":"Pagina / (home) — seccion de estadisticas y tarjetas de categorias"
      },
    ]
  },
  { "num":2,"name":"GESTION DE SESION","hbg":C_WINE_HX,"fgc":C_WINE,
    "desc":"Funcionalidades de autenticacion y gestion de identidad para usuarios registrados "
           "del sistema. Controlan el acceso a todas las funcionalidades protegidas.",
    "cases":[
      { "id":"CU-05","name":"Iniciar Sesion en el Sistema","priority":10,"scope":"OPERATIVO",
        "purpose":"Autenticar usuarios registrados para que accedan a funcionalidades exclusivas "
                   "segun su rol asignado (analista, gerente o administrador).",
        "desc":"El usuario ingresa username y password en el formulario /login. El sistema consulta "
               "la tabla usuarios_sistema en StarRocks, valida el hash bcrypt con check_password_hash "
               "y si es valido crea la sesion Flask con session['user'] y session['rol']. Redirige "
               "al dashboard del rol correspondiente. Registra el evento LOGIN en la tabla auditoria "
               "con IP, timestamp y usuario.",
        "stories":[
          ("Analista de datos","iniciar sesion con mis credenciales institucionales","acceder al dashboard analitico y comenzar mi jornada de analisis"),
          ("Gerente comercial","autenticarme en el sistema de forma segura","visualizar los reportes estrategicos y KPIs de ventas de vinos"),
          ("Administrador del sistema","iniciar sesion como administrador","gestionar usuarios del sistema y ejecutar procesos de carga de datos ETL"),
          ("Usuario registrado","que mis credenciales sean validadas con hash criptografico","garantizar que mi cuenta este protegida contra accesos no autorizados"),
        ],
        "img":"Pagina /login — formulario de inicio de sesion con validacion"
      },
      { "id":"CU-06","name":"Cerrar Sesion del Sistema","priority":9,"scope":"OPERATIVO",
        "purpose":"Terminar la sesion activa del usuario de forma segura, liberando los datos "
                   "de sesion y registrando el evento para trazabilidad.",
        "desc":"El usuario accede a /logout (GET o boton de navegacion). El sistema llama a "
               "session.clear() para destruir la sesion Flask, elimina las cookies de sesion "
               "del navegador y redirige al catalogo publico (/vinos). Registra el evento LOGOUT "
               "en la tabla auditoria con timestamp e IP del usuario.",
        "stories":[
          ("Analista","cerrar mi sesion al terminar mi turno de trabajo","evitar que otras personas accedan al sistema desde mi equipo sin autorizacion"),
          ("Gerente","que el sistema invalide completamente mi sesion","proteger la informacion confidencial si olvido cerrar el navegador"),
          ("Administrador","que el cierre de sesion quede registrado en el log","auditar todos los accesos y salidas del sistema para cumplimiento"),
          ("Usuario registrado","ser redirigido al catalogo publico tras cerrar sesion","seguir navegando el contenido publico sin necesidad de volver a iniciar sesion"),
        ],
        "img":"Navegacion con boton Cerrar Sesion y redireccion post-logout"
      },
    ]
  },
  { "num":3,"name":"ANALISIS DE DATOS","hbg":C_BLUE_HX,"fgc":C_BLUE,
    "desc":"Funcionalidades exclusivas para el rol Analista. Permiten explorar, filtrar y "
           "exportar los datos del dataset Winemag mediante consultas OLAP a StarRocks.",
    "cases":[
      { "id":"CU-07","name":"Visualizar Dashboard Analitico","priority":9,"scope":"TACTICO",
        "purpose":"Proporcionar al analista una vista interactiva con KPIs, graficos y tablas "
                   "derivadas del dataset Winemag para apoyar la toma de decisiones analiticas.",
        "desc":"El dashboard en /dashboard muestra: Top 10 paises por volumen de resenas, "
               "distribucion de precios por decil, matriz variedad vs. puntuacion promedio, "
               "y lista de bodegas destacadas. Los datos provienen de consultas OLAP directas "
               "a StarRocks via pymysql (puerto 9030). Los graficos usan Chart.js. "
               "Requiere sesion activa con rol analista.",
        "stories":[
          ("Analista de datos","ver el dashboard completo con todos los KPIs del sistema","identificar rapidamente tendencias y patrones en el mercado vinicola"),
          ("Analista","filtrar el dashboard por pais especifico","enfocar mi analisis en los mercados que manejo en mi area de responsabilidad"),
          ("Analista junior","ver la distribucion de precios en un grafico claro","detectar en que segmento de precio se concentra la mayor oferta del mercado"),
          ("Analista senior","que el dashboard se actualice con datos en tiempo real","trabajar siempre con la informacion mas actualizada disponible en el sistema"),
        ],
        "img":"Pagina /dashboard — dashboard analitico con graficos Chart.js y KPIs"
      },
      { "id":"CU-08","name":"Aplicar Filtros Avanzados de Analisis","priority":8,"scope":"TACTICO",
        "purpose":"Permitir al analista refinar las visualizaciones del dashboard mediante filtros "
                   "multi-criterio que se traducen en consultas OLAP dinamicas a StarRocks.",
        "desc":"El panel de filtros permite combinar: pais (dim_pais), variedad (dim_variedad), "
               "rango de precio (BETWEEN x AND y), rango de puntuacion (>= n), y texto de "
               "descripcion (LIKE). Al aplicar, el sistema construye la consulta SQL dinamicamente "
               "y actualiza todos los graficos del dashboard. Los filtros se mantienen en la sesion "
               "Flask hasta que el usuario los limpia.",
        "stories":[
          ("Analista","aplicar multiples filtros simultaneamente","cruzar dimensiones y encontrar patrones especificos imposibles de ver globalmente"),
          ("Analista de mercado","filtrar por rango de puntuacion y precio al mismo tiempo","analizar exclusivamente el segmento de vinos premium de alta calificacion"),
          ("Analista","que mis filtros persistan durante toda mi sesion activa","no tener que re-configurar los mismos filtros cada vez que navego entre paginas"),
          ("Analista senior","poder combinar hasta 5 criterios de filtrado diferentes","realizar analisis multidimensionales complejos sobre el dataset completo"),
        ],
        "img":"Panel lateral de filtros en /dashboard con multiples criterios activos"
      },
      { "id":"CU-09","name":"Generar Reporte de Analisis","priority":8,"scope":"TACTICO",
        "purpose":"Crear reportes estructurados con los datos y visualizaciones del analisis "
                   "actual, listos para presentar a stakeholders o archivar.",
        "desc":"El analista accede a /reportes, configura el reporte (titulo, descripcion, "
               "filtros a aplicar) y el sistema genera el documento con las tablas de datos, "
               "metricas calculadas y graficos del estado actual del dashboard. El reporte "
               "se almacena con timestamp y los filtros usados para garantizar trazabilidad. "
               "Puede exportarse como PDF o descargarse como CSV.",
        "stories":[
          ("Analista","generar un reporte con los resultados de mi analisis actual","presentar conclusiones fundamentadas en datos en la reunion semanal del equipo"),
          ("Analista","que el reporte incluya automaticamente los graficos del dashboard","tener una presentacion visual efectiva sin necesidad de capturas manuales"),
          ("Analista","poder darle nombre y descripcion personalizada a cada reporte","identificar facilmente cada reporte en el historial de reportes generados"),
          ("Analista","que el reporte incluya los filtros aplicados y la fecha de generacion","garantizar la trazabilidad y reproducibilidad del analisis documentado"),
        ],
        "img":"Pagina /reportes — formulario de configuracion y lista de reportes generados"
      },
      { "id":"CU-10","name":"Exportar Datos en CSV o PDF","priority":7,"scope":"TACTICO",
        "purpose":"Exportar los datos del analisis o reporte en formatos estandar para uso "
                   "en otras herramientas de analisis o para distribucion formal.",
        "desc":"El analista hace click en Exportar CSV (datos crudos con filtros aplicados) "
               "o Exportar PDF (reporte formateado). Para CSV: el sistema ejecuta la consulta "
               "OLAP con los filtros activos, serializa con csv.writer y retorna con "
               "Content-Disposition: attachment. Para PDF: usa reportlab o WeasyPrint para "
               "generar el documento con tablas y graficos.",
        "stories":[
          ("Analista","exportar los datos filtrados a formato CSV","procesarlos posteriormente en Excel, Python o Power BI para analisis adicionales"),
          ("Analista","generar un PDF del reporte de analisis","compartirlo con gerentes que no tienen acceso directo al sistema"),
          ("Analista","que el archivo CSV incluya todos los campos del esquema estrella","no perder ningun atributo importante en el proceso de exportacion"),
          ("Analista","que el PDF tenga formato profesional con logo y encabezado","presentarlo como documento oficial ante clientes o directivos"),
        ],
        "img":"Boton de exportacion en /reportes — dialogo de descarga de CSV o PDF"
      },
    ]
  },
  { "num":4,"name":"INTELIGENCIA GERENCIAL","hbg":C_GOLD_HX,"fgc":RGBColor(0x8A,0x60,0x10),
    "desc":"Funcionalidades exclusivas para el rol Gerente. Orientadas a la toma de decisiones "
           "estrategicas con vistas ejecutivas, comparacion de mercados y reportes de alto nivel.",
    "cases":[
      { "id":"CU-11","name":"Ver Dashboard Ejecutivo","priority":9,"scope":"ESTRATEGICO",
        "purpose":"Proporcionar al gerente una vista ejecutiva de alto nivel con los KPIs mas "
                   "relevantes para la toma de decisiones estrategicas sobre el negocio vinicola.",
        "desc":"La vista /gerente/dashboard muestra indicadores clave simplificados: volumen "
               "de resenas por region, precio promedio por mercado, puntuacion media por pais, "
               "y top 3 variedades del periodo. El diseno es limpio y visualmente impactante, "
               "disenado para lectura rapida en reuniones de directorio sin necesidad de "
               "conocimientos tecnicos. Requiere rol gerente en la sesion.",
        "stories":[
          ("Gerente general","ver todos los KPIs principales del negocio en una sola pantalla","tomar decisiones estrategicas informadas en menos de 2 minutos"),
          ("Gerente comercial","ver el performance de ventas por mercado geografico","identificar rapidamente en que paises enfocar los esfuerzos comerciales"),
          ("Director de area","que el dashboard sea visualmente limpio y profesional","presentarlo directamente en reuniones de directorio sin preparacion adicional"),
          ("Gerente regional","filtrar el dashboard ejecutivo para ver solo mi region","monitorear especificamente el desempeno del mercado bajo mi responsabilidad"),
        ],
        "img":"Pagina /gerente/dashboard — vista ejecutiva con KPIs y graficos de alto nivel"
      },
      { "id":"CU-12","name":"Comparar KPIs por Mercado/Pais","priority":8,"scope":"ESTRATEGICO",
        "purpose":"Facilitar al gerente la comparacion directa de indicadores de desempeno entre "
                   "diferentes mercados para identificar oportunidades de expansion o brechas.",
        "desc":"El gerente selecciona 2 a 4 paises o regiones desde un selector multiple. "
               "El sistema genera consultas OLAP paralelas a StarRocks para cada mercado "
               "y presenta una vista comparativa con graficos de barras lado a lado, tabla "
               "de diferencias absolutas y porcentuales, y ranking combinado por indice "
               "calidad-precio.",
        "stories":[
          ("Gerente de expansion","comparar el volumen de resenas entre Chile, Argentina y Espana","decidir en que mercado latinoamericano conviene enfocar la estrategia de inversion"),
          ("Director de ventas","ver el precio promedio comparado entre los top 5 paises","ajustar nuestra estrategia de pricing segun el posicionamiento de cada mercado"),
          ("Gerente de producto","comparar la puntuacion promedio entre diferentes regiones","identificar cuales son los mercados con mayor reputacion de calidad vinicola"),
          ("CEO","obtener un ranking de paises por indice combinado de precio y calidad","priorizar los mercados estrategicos para la expansion internacional del negocio"),
        ],
        "img":"Pagina /gerente/comparar — vista comparativa de KPIs entre mercados seleccionados"
      },
      { "id":"CU-13","name":"Generar Reporte Estrategico","priority":7,"scope":"ESTRATEGICO",
        "purpose":"Crear reportes ejecutivos formateados con narrativa, KPIs seleccionados y "
                   "visualizaciones para distribucion a nivel directivo y toma de decisiones.",
        "desc":"El gerente selecciona el periodo de analisis, los KPIs a incluir y el mercado "
               "de enfoque. El sistema genera un PDF con formato ejecutivo que incluye: "
               "resumen ejecutivo automatico, graficos de tendencias, tabla de desempeno "
               "comparativa y conclusiones con los datos mas relevantes del periodo.",
        "stories":[
          ("Gerente","generar un reporte mensual en PDF con formato ejecutivo","presentar el performance vinicola del mes al directorio de la empresa"),
          ("Gerente comercial","que el reporte incluya graficos de tendencias automaticamente","tener un documento completo sin necesidad de construirlo manualmente"),
          ("Gerente","poder seleccionar cuales KPIs incluir en cada reporte","personalizar el documento segun la audiencia y el objetivo de la presentacion"),
          ("CFO","obtener un reporte con comparativo del mes actual vs el anterior","evaluar la tendencia del negocio y tomar decisiones de presupuesto informadas"),
        ],
        "img":"Pagina /gerente/reportes — configurador de reporte estrategico y descarga en PDF"
      },
      { "id":"CU-14","name":"Analizar Tendencias de Precios","priority":7,"scope":"ESTRATEGICO",
        "purpose":"Visualizar la evolucion de los precios del mercado vinicola para identificar "
                   "tendencias y fundamentar decisiones de pricing y portafolio.",
        "desc":"El sistema presenta graficos de lineas con la evolucion de precios por variedad, "
               "pais o region en el periodo seleccionado. Las consultas agrupan por anio "
               "extrayendo el anio del titulo del vino. Incluye linea de tendencia calculada "
               "con regresion lineal simple y proyeccion para los proximos 2 periodos.",
        "stories":[
          ("Gerente de compras","ver la tendencia historica de precios de las variedades principales","negociar mejor con proveedores usando datos concretos del mercado"),
          ("Analista de mercado","identificar variedades con precio en tendencia alcista","recomendar inversiones en etiquetas con mayor potencial de revalorizacion"),
          ("Gerente de portafolio","comparar la evolucion de precios entre dos variedades clave","decidir cuales incluir o retirar del portafolio comercial"),
          ("Director financiero","ver una proyeccion de precios para el proximo trimestre","planificar el presupuesto de compras de forma fundamentada en tendencias"),
        ],
        "img":"Pagina /gerente/tendencias — grafico de lineas con tendencias de precios por variedad"
      },
    ]
  },
  { "num":5,"name":"ADMINISTRACION DEL SISTEMA","hbg":C_PURPLE_HX,"fgc":C_PURPLE,
    "desc":"Funcionalidades exclusivas para el rol Administrador. Cubren la gestion completa "
           "del sistema: usuarios, datos, respaldos, auditoria y monitoreo de infraestructura.",
    "cases":[
      { "id":"CU-15","name":"Gestionar Usuarios del Sistema","priority":10,"scope":"OPERATIVO",
        "purpose":"Permitir al administrador crear, modificar y eliminar cuentas de usuario "
                   "con asignacion de roles y control de estado de cuenta.",
        "desc":"En /admin/usuarios el admin ve la lista de todos los usuarios con su rol (admin/"
               "analista/gerente) y estado (activo/inactivo). Puede crear nuevos usuarios con "
               "hash bcrypt de la contrasena, editar el rol o estado, y eliminar cuentas. "
               "Todas las operaciones CRUD quedan registradas en la tabla auditoria con "
               "el tipo de accion CRUD_USER.",
        "stories":[
          ("Administrador","crear nuevos usuarios con roles especificos asignados","dar acceso controlado y seguro al sistema a nuevos empleados de la organizacion"),
          ("Administrador","desactivar la cuenta de un usuario sin eliminarla","revocar acceso a empleados que ya no trabajan sin perder el historial de sus acciones"),
          ("Administrador","cambiar el rol de un usuario existente","ajustar sus permisos y acceso cuando cambia de posicion dentro de la empresa"),
          ("Administrador","ver el historial de acciones de cada usuario en el sistema","auditar el uso del sistema y detectar comportamientos anomalos o no autorizados"),
        ],
        "img":"Pagina /admin/usuarios — tabla de usuarios con opciones CRUD y asignacion de roles"
      },
      { "id":"CU-16","name":"Ejecutar Proceso ETL","priority":10,"scope":"OPERATIVO",
        "purpose":"Cargar y actualizar el dataset Winemag en StarRocks mediante el pipeline "
                   "ETL (Extract-Transform-Load) administrado directamente desde el sistema.",
        "desc":"El admin accede a /admin/etl y puede: (1) Subir un nuevo CSV Winemag, "
               "(2) Iniciar el proceso ETL. El pipeline ejecuta: Extraccion del CSV con pandas, "
               "Transformacion (normalizar dimensiones, calcular FKs, limpiar nulos), y Carga "
               "en StarRocks via STREAM LOAD o pymysql bulk insert. El sistema muestra el "
               "progreso en tiempo real via Server-Sent Events y registra ETL_START/ETL_END "
               "en auditoria.",
        "stories":[
          ("Administrador","ejecutar el proceso ETL con un dataset CSV actualizado","cargar las ultimas resenas de vinos y mantener el catalogo vigente"),
          ("Administrador","ver el progreso del proceso ETL en tiempo real","saber exactamente en que etapa esta la carga y si esta funcionando correctamente"),
          ("Administrador","que el ETL valide los datos antes de insertarlos en la BD","evitar datos corruptos, duplicados o malformados en la base de datos de produccion"),
          ("Administrador","que el ETL reporte errores detallados por fila","diagnosticar y corregir rapidamente problemas especificos en el archivo de carga"),
        ],
        "img":"Pagina /admin/etl — panel de carga ETL con progreso en tiempo real"
      },
      { "id":"CU-17","name":"Gestionar Respaldos del Sistema","priority":8,"scope":"OPERATIVO",
        "purpose":"Crear, programar y restaurar respaldos de la base de datos y configuraciones "
                   "del sistema para garantizar la continuidad operacional del negocio.",
        "desc":"En /admin/respaldos el admin puede crear respaldos manuales del esquema "
               "operacional (usuarios_sistema + auditoria), descargarlos como archivos SQL "
               "comprimidos, y restaurar desde un archivo previo. Tambien configura respaldos "
               "automaticos programados. Cada operacion queda en auditoria con tipo BACKUP o "
               "RESTORE.",
        "stories":[
          ("Administrador","crear respaldos manuales del sistema en cualquier momento","asegurar una copia de seguridad antes de realizar cambios importantes en el sistema"),
          ("Administrador","programar respaldos automaticos nocturnos","garantizar respaldos diarios sin necesidad de intervension manual repetitiva"),
          ("Administrador","restaurar un respaldo especifico desde el historial","recuperar el sistema completamente en caso de fallo o corrupcion de datos"),
          ("Gerente de TI","que los respaldos incluyan todos los datos operacionales criticos","garantizar la recuperacion completa del sistema ante cualquier escenario de desastre"),
        ],
        "img":"Pagina /admin/respaldos — historial de respaldos con opciones de descarga y restauracion"
      },
      { "id":"CU-18","name":"Consultar Log de Auditoria","priority":9,"scope":"OPERATIVO",
        "purpose":"Proporcionar al administrador visibilidad completa sobre todas las acciones "
                   "realizadas en el sistema para garantizar trazabilidad, seguridad y cumplimiento.",
        "desc":"En /admin/auditoria el admin consulta la tabla auditoria de StarRocks. Puede "
               "filtrar por: usuario, tipo de accion (LOGIN/LOGOUT/ETL/BACKUP/CRUD_USER/EXPORT), "
               "rango de fechas e IP. La tabla muestra: timestamp, usuario, rol, accion, detalle "
               "y IP. Los registros son inmutables (DUPLICATE KEY, solo INSERT). Exportable a CSV.",
        "stories":[
          ("Administrador","ver todas las acciones de todos los usuarios en el sistema","detectar comportamientos inusuales o accesos no autorizados de forma proactiva"),
          ("Auditor interno","filtrar el log de auditoria por un usuario especifico","revisar completamente todas las acciones realizadas por un empleado en particular"),
          ("Administrador de seguridad","ver los intentos de inicio de sesion fallidos","identificar posibles ataques de fuerza bruta o intentos de acceso no autorizado"),
          ("Gerente de cumplimiento","exportar el log completo de auditoria en CSV","cumplir con los requisitos de compliance de normativas y auditorias externas"),
        ],
        "img":"Pagina /admin/auditoria — tabla de log con filtros por usuario, accion y fecha"
      },
      { "id":"CU-19","name":"Monitorear Estado del Sistema","priority":7,"scope":"OPERATIVO",
        "purpose":"Proporcionar al administrador una vista del estado de salud de todos los "
                   "servicios del sistema y metricas de rendimiento operacional.",
        "desc":"En /admin/monitor el admin ve el estado de cada servicio Docker (Flask :5000, "
               "StarRocks FE :9030, PocketBase :8090), conectividad a la base de datos, "
               "espacio en disco disponible, y estado de los ultimos procesos ETL y respaldos. "
               "Las alertas visuales (rojo/verde) indican si algun componente esta caido. "
               "Los datos se actualizan via polling AJAX cada 30 segundos.",
        "stories":[
          ("Administrador","ver el estado de todos los servicios del sistema en tiempo real","identificar rapidamente si algun componente esta caido y actuar de inmediato"),
          ("Administrador","recibir alerta visual clara cuando StarRocks no esta disponible","minimizar el tiempo de inactividad del sistema tomando accion inmediata"),
          ("Administrador de infraestructura","ver el uso actual de disco y memoria del servidor","planificar el escalado del sistema antes de que los recursos se agoten completamente"),
          ("Gerente de TI","ver cuando fue ejecutado el ultimo ETL y el ultimo respaldo","garantizar que todos los procesos automatizados del sistema estan funcionando correctamente"),
        ],
        "img":"Pagina /admin/monitor — panel de estado de servicios con indicadores de salud"
      },
    ]
  },
]

def render_packages(doc):
    doc.add_page_break()
    heading_styled(doc, "3. DESCRIPCION DE LOS CASOS DE USO AGRUPADOS POR PAQUETES",
                   size=15, color=C_WINED, space_before=0)
    para(doc,
         "Los casos de uso se agrupan en 5 paquetes segun el actor principal y el nivel "
         "de acceso requerido. Cada paquete agrupa funcionalidades cohesivas del sistema.",
         size=11, space_after=12)

    for pkg in PACKAGES:
        doc.add_page_break()
        # ── HEADER DEL PAQUETE ─────────────────────────────────────────────
        t = doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        shd(cell, pkg["hbg"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"PAQUETE {pkg['num']}: {pkg['name']}")
        r.bold=True; r.font.size=Pt(14); r.font.color.rgb=C_WHITE
        doc.add_paragraph()
        para(doc, pkg["desc"], size=11, space_after=10)

        for cu in pkg["cases"]:
            # separador visual
            t2 = doc.add_table(rows=1, cols=1)
            cell2 = t2.rows[0].cells[0]
            shd(cell2, pkg["hbg"])
            p2 = cell2.paragraphs[0]
            r2 = p2.add_run(f"{cu['id']}  |  {cu['name']}")
            r2.bold=True; r2.font.size=Pt(13); r2.font.color.rgb=C_WHITE
            doc.add_paragraph()

            # PRIORIDAD
            p3 = doc.add_paragraph()
            p3.paragraph_format.space_before=Pt(0); p3.paragraph_format.space_after=Pt(4)
            r_lbl = p3.add_run("PRIORIDAD: ")
            r_lbl.bold=True; r_lbl.font.size=Pt(11)
            r_val = p3.add_run(f"{cu['priority']}/10  [{cu['scope']}]")
            r_val.font.size=Pt(11); r_val.font.color.rgb=pkg["fgc"]

            # PROPOSITO
            p4 = doc.add_paragraph()
            p4.paragraph_format.space_before=Pt(0); p4.paragraph_format.space_after=Pt(4)
            r4l = p4.add_run("PROPOSITO DEL CU: ")
            r4l.bold=True; r4l.font.size=Pt(11)
            p4.add_run(cu["purpose"]).font.size=Pt(11)

            # DESCRIPCION
            p5 = doc.add_paragraph()
            p5.paragraph_format.space_before=Pt(4); p5.paragraph_format.space_after=Pt(4)
            r5l = p5.add_run("DESCRIPCION: ")
            r5l.bold=True; r5l.font.size=Pt(11)
            p5.add_run(cu["desc"]).font.size=Pt(11)

            # HISTORIAS DE USUARIO
            p6 = doc.add_paragraph()
            p6.paragraph_format.space_before=Pt(6); p6.paragraph_format.space_after=Pt(2)
            r6 = p6.add_run("HISTORIAS DE USUARIO:")
            r6.bold=True; r6.font.size=Pt(11)

            t3 = doc.add_table(rows=1+len(cu["stories"]), cols=3)
            set_borders(t3)
            for ci,(h,bh) in enumerate([("ROL",pkg["hbg"]),("DESEO",pkg["hbg"]),("PROPOSITO",pkg["hbg"])]):
                ch = t3.rows[0].cells[ci]; shd(ch, bh)
                rh = ch.paragraphs[0].add_run(h)
                rh.bold=True; rh.font.color.rgb=C_WHITE; rh.font.size=Pt(10)
            for si,(rol,deseo,prop) in enumerate(cu["stories"]):
                bg_s = C_GRAYL_HX if si%2==0 else C_WHITE_HX
                row_s = t3.rows[si+1].cells
                for c in row_s: shd(c, bg_s)
                row_s[0].paragraphs[0].add_run(f"Como {rol}").font.size=Pt(10)
                row_s[1].paragraphs[0].add_run(f"deseo {deseo}").font.size=Pt(10)
                row_s[2].paragraphs[0].add_run(f"para {prop}").font.size=Pt(10)
            doc.add_paragraph()

            # IMAGEN
            p7 = doc.add_paragraph()
            p7.paragraph_format.space_before=Pt(4); p7.paragraph_format.space_after=Pt(2)
            r7 = p7.add_run("IMAGEN DEL CASO DE USO IMPLEMENTADO EN EL SISTEMA:")
            r7.bold=True; r7.font.size=Pt(11)
            placeholder_img(doc, cu["img"])
            doc.add_paragraph()

# ── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── portada ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    heading_styled(doc, "VINANALYTICS GROUP", size=26, color=C_WINED, align="center",
                   space_before=60, space_after=4)
    heading_styled(doc, "Sistema de Analisis de Vinos Finos", size=16, color=C_GRAY,
                   align="center", bold=False, space_before=0, space_after=30)

    t_cov = doc.add_table(rows=1, cols=1)
    shd(t_cov.rows[0].cells[0], C_WINEL_HX)
    p_cov = t_cov.rows[0].cells[0].paragraphs[0]
    p_cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cov = p_cov.add_run(
        "Documentacion del Sistema  |  Diagramas y Casos de Uso\n"
        "Stack: Flask · StarRocks OLAP · PocketBase · Docker\n"
        "Dataset: Winemag 308,724 resenas | 44 paises | 708 variedades"
    )
    r_cov.font.size=Pt(11); r_cov.font.color.rgb=C_WINED

    para(doc,"", space_before=30, space_after=0)
    para(doc, "Universidad Tecnica Estatal de Quevedo  —  Ingenieria en Sistemas",
         align="center", size=11, color=C_GRAY)
    para(doc, "2026", align="center", size=11, color=C_GRAY)

    # ── secciones ───────────────────────────────────────────────────────────
    section_db(doc)
    section_uc_diagram(doc)
    render_packages(doc)

    # save
    out = os.path.join(OUT, "VinAnalytics_Documentacion.docx")
    doc.save(out)
    print(f"OK  {out}")

if __name__ == "__main__":
    main()
