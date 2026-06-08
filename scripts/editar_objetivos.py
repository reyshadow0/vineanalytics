"""Edita Objetivos.docx con el contenido completo de VinAnalytics."""
import sys, os
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colores ──────────────────────────────────────────────────────────────────
C_WINED  = RGBColor(0x5C,0x0E,0x0E);  HX_WINEL  = "FFEAEA"; HX_WINE  = "8B1A1A"
C_BLUE   = RGBColor(0x1A,0x4A,0x8A);  HX_BLUEL  = "E8F0FF"; HX_BLUE  = "1A4A8A"
C_GREEN  = RGBColor(0x2A,0x6A,0x10);  HX_GREENL = "EAFAEE"; HX_GREEN = "2A6A10"
C_PURPLE = RGBColor(0x6A,0x1A,0x8A);  HX_PURPL  = "F5E8FF"
C_GOLD   = RGBColor(0x8A,0x60,0x10);  HX_GOLDL  = "FFF8E8"
C_WHITE  = RGBColor(0xFF,0xFF,0xFF);  HX_WHITE  = "FFFFFF"
C_GRAY   = RGBColor(0x55,0x55,0x55);  HX_GRAYL  = "F5F5F5"

# ── helpers ───────────────────────────────────────────────────────────────────
def shd(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(s)

def valign(cell, val="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)

def write_cell(cell, text, bold=False, size=10, color=None, bg=None, align="left"):
    tc = cell._tc
    for p_el in list(tc.findall(qn("w:p"))):
        tc.remove(p_el)
    p_el = OxmlElement("w:p")
    tc.append(p_el)
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bg:
        shd(cell, bg)

def set_para(p, text, size=11, color=None):
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
        p.runs[0].font.size = Pt(size)
        if color:
            p.runs[0].font.color.rgb = color
    else:
        r = p.add_run(text)
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color

# ── DATOS VINANALYTICS ────────────────────────────────────────────────────────
DESCRIPCION = (
    "VinAnalytics Group es una plataforma de inteligencia de mercado vinícola que "
    "centraliza, procesa y visualiza datos de reseñas de vinos a nivel global. "
    "Desarrollada con Flask, StarRocks (motor OLAP columnar), PocketBase y Docker, "
    "el sistema gestiona más de 308,000 reseñas de vinos provenientes de 44 países "
    "y 708 variedades de uva, procesadas a través de un pipeline ETL automatizado. "
    "Está orientada a analistas de datos, gerentes comerciales y administradores del "
    "sector vitivinícola, permitiendo análisis estadístico en tiempo real, generación "
    "de reportes estratégicos y gestión integral del catálogo vinícola."
)

VISION = (
    "Ser la plataforma líder de análisis de inteligencia de mercado vinícola en "
    "América Latina, proporcionando información estratégica basada en datos que "
    "impulse decisiones competitivas y sostenibles en la industria del vino."
)

MISION = (
    "Transformar datos masivos de reseñas de vinos en información accionable para "
    "empresas del sector vitivinícola, mediante un sistema analítico de alto rendimiento "
    "que democratiza el acceso a la inteligencia de mercado, apoya la toma de decisiones "
    "basada en datos y garantiza la seguridad y trazabilidad de la información."
)

# Jerarquía OE/OT/OO/METAS
# Cada entrada: (oe_texto, ot_texto, oo_texto, meta_texto)
# None = celda que se fusionará (pertenece al mismo OE u OT del anterior)
HIERARCHY = [
    # ── OE1 (filas 1-4) ────────────────────────────────────────────────────
    ("OE1: Mejorar la toma de decisiones comerciales mediante inteligencia de datos vinícola",
     "OT1.1: Proveer análisis del mercado vinícola en tiempo real a los analistas",
     "OO1.1.1: Visualizar KPIs del mercado en un dashboard interactivo con filtros dinámicos",
     "Reducir en un 40% el tiempo destinado al análisis manual de datos de mercado"),
    (None, None,
     "OO1.1.2: Generar reportes de análisis personalizados por filtros, variables y períodos",
     "Producir al menos 10 reportes analíticos mensuales por usuario analista activo"),
    (None,
     "OT1.2: Facilitar la comparación estratégica entre mercados y países vinícolas",
     "OO1.2.1: Comparar KPIs de hasta 4 mercados vinícolas de forma simultánea",
     "Reducir el tiempo de comparación entre mercados de horas a menos de 5 minutos"),
    (None, None,
     "OO1.2.2: Analizar tendencias de precios por variedad con proyección futura",
     "Proyectar tendencias de precio con base histórica para la planificación trimestral"),
    # ── OE2 (filas 5-8) ────────────────────────────────────────────────────
    ("OE2: Optimizar la gestión y seguridad de la información vinícola del sistema",
     "OT2.1: Centralizar y mantener actualizado el catálogo de reseñas de vinos",
     "OO2.1.1: Ejecutar el pipeline ETL para la carga masiva del dataset Winemag",
     "Cargar más de 300,000 registros en menos de 10 minutos con 0% de errores críticos"),
    (None, None,
     "OO2.1.2: Garantizar acceso continuo y eficiente al catálogo para todos los usuarios",
     "Mantener disponibilidad del catálogo 24/7 con tiempo de respuesta menor a 2 segundos"),
    (None,
     "OT2.2: Garantizar la seguridad, integridad y trazabilidad de los datos del sistema",
     "OO2.2.1: Gestionar usuarios del sistema con control de acceso basado en roles (RBAC)",
     "Asegurar que el 100% de los accesos al sistema estén validados por rol asignado"),
    (None, None,
     "OO2.2.2: Registrar y auditar todas las acciones realizadas por usuarios en el sistema",
     "Mantener trazabilidad del 100% de las operaciones con registro inmutable en auditoría"),
    # ── OE3 (filas 9-10) ───────────────────────────────────────────────────
    ("OE3: Mejorar la experiencia de acceso a información vinícola para todos los usuarios",
     "OT3.1: Facilitar el acceso público y sin barreras al catálogo de vinos del sistema",
     "OO3.1.1: Permitir la navegación libre del catálogo sin necesidad de registro previo",
     "Atender visitantes sin registro con tiempo de carga de página menor a 2 segundos"),
    (None, None,
     "OO3.1.2: Proveer búsqueda y filtrado avanzado multi-criterio sobre todo el catálogo",
     "Retornar resultados de búsqueda filtrada en menos de 1 segundo de tiempo de respuesta"),
]

# Tabla 2 — aporte del sistema a cada objetivo (reemplaza los 3 ejemplos genéricos)
T2_ROWS = [
    ("OE1: Mejorar la toma de decisiones comerciales",
     "OT1.1: Proveer análisis de mercado en tiempo real",
     "OO1.1.1: Visualizar KPIs del mercado en dashboard",
     "Análisis OLAP en tiempo real",
     "El sistema ejecuta consultas OLAP a StarRocks y presenta KPIs del mercado vinícola "
     "en gráficos interactivos (Chart.js): top países por volumen de reseñas, distribución "
     "de precios y puntuación promedio. Permite filtros por país, variedad y período.",
     "% de reducción en tiempo de análisis manual. Meta: 40%.\nN° de consultas por sesión.",
     "Actor: Analista\nCU-07: Visualizar Dashboard Analítico"),
    ("OE1: Mejorar la toma de decisiones comerciales",
     "OT1.1: Proveer análisis de mercado en tiempo real",
     "OO1.1.2: Generar reportes de análisis personalizados",
     "Generación de Reportes Analíticos",
     "El sistema permite al analista configurar y generar reportes con los datos del análisis "
     "actual (filtros aplicados, métricas calculadas, gráficos), exportables en CSV o PDF "
     "con encabezado corporativo y trazabilidad de parámetros usados.",
     "N° de reportes generados por mes. Meta: ≥10/analista/mes.\nFormatos disponibles: CSV, PDF.",
     "Actor: Analista\nCU-09: Generar Reporte de Análisis\nCU-10: Exportar Datos CSV/PDF"),
    ("OE1: Mejorar la toma de decisiones comerciales",
     "OT1.2: Facilitar comparación estratégica entre mercados",
     "OO1.2.1: Comparar KPIs entre múltiples mercados simultáneamente",
     "Inteligencia Gerencial Comparativa",
     "El sistema genera vistas comparativas entre 2 a 4 mercados seleccionados por el gerente: "
     "gráficos lado a lado, tabla de diferencias absolutas y porcentuales, y ranking combinado "
     "calidad-precio. Consultas OLAP paralelas a StarRocks.",
     "N° de mercados analizados simultáneamente. Meta: hasta 4.\nTiempo de comparación: <30 s.",
     "Actor: Gerente\nCU-12: Comparar KPIs por Mercado/País"),
    ("OE1: Mejorar la toma de decisiones comerciales",
     "OT1.2: Facilitar comparación estratégica entre mercados",
     "OO1.2.2: Analizar tendencias de precios por variedad",
     "Análisis de Tendencias de Precios",
     "El sistema presenta gráficos de líneas con evolución de precios por variedad, país o región "
     "en el período seleccionado. Incluye línea de tendencia con regresión lineal y proyección "
     "para los próximos 2 períodos.",
     "% de precisión en proyección de precios. Meta: ≥85%.\nPeríodo de proyección: 2 trimestres.",
     "Actor: Gerente\nCU-14: Analizar Tendencias de Precios"),
    ("OE2: Optimizar la gestión de información vinícola",
     "OT2.1: Centralizar y actualizar el catálogo de reseñas",
     "OO2.1.1: Ejecutar pipeline ETL del dataset Winemag",
     "Carga de Datos ETL (Extract-Transform-Load)",
     "El sistema ejecuta el pipeline ETL: extrae el CSV Winemag con pandas, normaliza las 6 "
     "dimensiones (país, variedad, bodega, catador, región, provincia), calcula claves foráneas "
     "y carga en StarRocks vía STREAM LOAD con reporte de progreso en tiempo real.",
     "Tiempo de carga. Meta: <10 min para 300K+ registros.\nTasa de error crítico: 0%.",
     "Actor: Administrador\nCU-16: Ejecutar Proceso ETL"),
    ("OE2: Optimizar la gestión de información vinícola",
     "OT2.2: Garantizar seguridad e integridad de datos",
     "OO2.2.1: Gestionar usuarios con control RBAC",
     "Gestión de Acceso por Roles (RBAC)",
     "El sistema gestiona cuentas con tres roles diferenciados (analista, gerente, admin). "
     "Valida contraseñas con hash bcrypt, mantiene sesiones Flask y restringe cada rol "
     "únicamente a las rutas autorizadas. CRUD completo con auditoría de cambios.",
     "% de accesos controlados por rol. Meta: 100% RBAC.\nN° de usuarios activos/inactivos.",
     "Actor: Administrador\nCU-15: Gestionar Usuarios del Sistema"),
    ("OE2: Optimizar la gestión de información vinícola",
     "OT2.2: Garantizar seguridad e integridad de datos",
     "OO2.2.2: Auditar todas las acciones del sistema",
     "Log de Auditoría Inmutable",
     "El sistema registra en StarRocks (tabla auditoria, DUPLICATE KEY) cada acción realizada: "
     "LOGIN, LOGOUT, ETL, BACKUP, CRUD_USER, EXPORT. Campos: timestamp, usuario, rol, "
     "acción, detalle e IP. Filtrable por usuario, acción y fecha. Exportable en CSV.",
     "% de operaciones auditadas. Meta: 100%.\nN° de eventos registrados por día.",
     "Actor: Administrador\nCU-18: Consultar Log de Auditoría"),
    ("OE3: Mejorar la experiencia de acceso a información vinícola",
     "OT3.1: Facilitar acceso público sin barreras al catálogo",
     "OO3.1.1: Navegación libre del catálogo sin registro previo",
     "Catálogo de Vinos Público",
     "El sistema permite navegar el catálogo de 308,724 reseñas sin autenticación. "
     "Muestra tarjetas con foto real (Unsplash CDN), puntuación, precio y bodega. "
     "Paginación de 20 vinos por página con respuesta dinámica de StarRocks.",
     "Tiempo de carga de página. Meta: <2 segundos.\nN° de visitantes únicos por mes.",
     "Actor: Visitante Público\nCU-01: Explorar Catálogo de Vinos"),
    ("OE3: Mejorar la experiencia de acceso a información vinícola",
     "OT3.1: Facilitar acceso público sin barreras al catálogo",
     "OO3.1.2: Búsqueda y filtrado multi-criterio del catálogo",
     "Búsqueda y Filtrado Avanzado",
     "El sistema ofrece filtros combinables por país, variedad, rango de precio y puntuación. "
     "Las consultas se ejecutan en tiempo real contra StarRocks retornando resultados paginados. "
     "Los filtros son acumulables y se transmiten como parámetros GET.",
     "Tiempo de respuesta de búsqueda. Meta: <1 segundo.\nN° de filtros combinables disponibles: 5.",
     "Actor: Visitante Público\nCU-02: Buscar y Filtrar Vinos"),
]

# ── MAIN ──────────────────────────────────────────────────────────────────────
def main():
    path = os.path.join(os.path.dirname(__file__), "..", "diagramas", "Objetivos.docx")
    doc = Document(path)
    paras = doc.paragraphs

    # 1. Párrafos de texto libre ─────────────────────────────────────────────
    set_para(paras[1], DESCRIPCION)
    set_para(paras[4], VISION)
    set_para(paras[6], MISION)

    # 2. TABLA 1 — Jerarquía OE / OT / OO / METAS ───────────────────────────
    t1 = doc.tables[1]

    # Borrar todas las filas de datos (conservar solo cabecera)
    while len(t1.rows) > 1:
        t1._tbl.remove(t1.rows[-1]._tr)

    # Agregar 10 filas limpias
    current_oe, current_ot = "", ""
    for oe, ot, oo, meta in HIERARCHY:
        new_row = t1.add_row()
        if oe: current_oe = oe
        if ot: current_ot = ot
        oe_display = current_oe if oe else f"(cont.) {current_oe[:50]}"
        ot_display = current_ot if ot else f"(cont.) {current_ot[:50]}"
        write_cell(new_row.cells[0], oe_display, bold=bool(oe), size=9,
                   color=C_WINED, bg=HX_WINEL)
        write_cell(new_row.cells[1], ot_display, bold=bool(ot), size=9,
                   color=C_BLUE,  bg=HX_BLUEL)
        write_cell(new_row.cells[2], oo,   bold=False, size=10, bg=HX_GRAYL)
        write_cell(new_row.cells[3], meta, bold=False, size=10, color=C_GREEN,
                   bg=HX_GREENL)

    # 3. TABLA 2 — Aporte del sistema a los objetivos ─────────────────────────
    t2 = doc.tables[2]

    # Borrar todas las filas de datos (conservar solo cabecera)
    while len(t2.rows) > 1:
        t2._tbl.remove(t2.rows[-1]._tr)

    # Agregar todas las filas de VinAnalytics limpias
    for ri, row_data in enumerate(T2_ROWS):
        new_row = t2.add_row()
        bg = HX_GRAYL if ri % 2 == 0 else HX_WHITE
        for ci, text in enumerate(row_data):
            write_cell(new_row.cells[ci], text, size=9, bg=bg)

    doc.save(path)
    print("OK  Objetivos.docx editado —", len(HIERARCHY), "objetivos,", len(T2_ROWS), "filas de aporte del sistema")

if __name__ == "__main__":
    main()
